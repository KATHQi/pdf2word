# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import time
import subprocess
import re
from copy import deepcopy

from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger
logger = get_logger()


def group_text_by_lines(text_results, y_threshold=10):
    """
    根据文本框的y坐标将文本分组到不同的行
    相近y坐标的文本框被认为是同一行
    
    Args:
        text_results: 包含 {'text': str, 'text_region': [[x1,y1], [x2,y2], [x3,y3], [x4,y4]] 或 [x1,y1,x2,y2]} 的列表
        y_threshold: y坐标差异阈值，小于此值认为是同一行
    
    Returns:
        按行分组的文本列表，每行内按x坐标排序
    """
    if not text_results:
        return []
    
    # 提取每个文本框的位置信息
    items = []
    for item in text_results:
        if not isinstance(item, dict):
            continue
        
        text = item.get('text', '')
        text_region = item.get('text_region', None)
        
        if text_region is None:
            # 如果没有 text_region，每个文本单独成行
            items.append({'text': text, 'x': 0, 'y': len(items) * 100, 'x2': 100})
            continue
        
        # 处理不同格式的 text_region
        if isinstance(text_region, list):
            if len(text_region) == 4 and isinstance(text_region[0], (int, float)):
                # 格式: [x1, y1, x2, y2]
                x1, y1, x2, y2 = text_region
            elif len(text_region) >= 4 and isinstance(text_region[0], list):
                # 格式: [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
                x1 = text_region[0][0]
                y1 = text_region[0][1]
                x2 = text_region[2][0]
                y2 = text_region[2][1]
            else:
                # 格式不对，单独成行
                items.append({'text': text, 'x': 0, 'y': len(items) * 100, 'x2': 100})
                continue
        else:
            items.append({'text': text, 'x': 0, 'y': len(items) * 100, 'x2': 100})
            continue
        
        # 使用文本框中心y作为行判断依据
        center_y = (y1 + y2) / 2 if y2 > y1 else y1
        
        items.append({
            'text': text,
            'x': x1,
            'y': center_y,
            'x2': x2
        })
    
    if not items:
        return []
    
    # 按y坐标排序
    items.sort(key=lambda x: (x['y'], x['x']))
    
    # 分组到行
    lines = []
    current_line = [items[0]]
    current_y = items[0]['y']
    
    for item in items[1:]:
        if abs(item['y'] - current_y) <= y_threshold:
            # 同一行
            current_line.append(item)
        else:
            # 新行
            # 当前行按x排序
            current_line.sort(key=lambda x: x['x'])
            lines.append(current_line)
            current_line = [item]
            current_y = item['y']
    
    # 添加最后一行
    if current_line:
        current_line.sort(key=lambda x: x['x'])
        lines.append(current_line)
    
    return lines


def merge_line_texts(line_items, use_spacing=True, char_width=10):
    """
    合并同一行的文本，保持适当的间距
    
    Args:
        line_items: 同一行的文本项列表
        use_spacing: 是否根据x坐标添加空格
        char_width: 估算的每个字符宽度（像素）
    
    Returns:
        合并后的行文本
    """
    if not line_items:
        return ""
    
    if len(line_items) == 1:
        return line_items[0]['text']
    
    result_parts = []
    for i, item in enumerate(line_items):
        if i == 0:
            result_parts.append(item['text'])
        else:
            # 计算与前一个文本框的间距
            gap = item['x'] - line_items[i-1]['x2']
            
            if use_spacing and gap > char_width * 2:
                # 较大间距，添加多个空格或制表符
                num_spaces = max(1, int(gap / char_width / 2))
                result_parts.append(' ' * num_spaces)
            else:
                # 普通间距，添加一个空格
                result_parts.append(' ')
            
            result_parts.append(item['text'])
    
    return ''.join(result_parts)


def _escape_html(text):
    """转义HTML特殊字符"""
    if not isinstance(text, str):
        text = str(text)
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&#39;')
    return text


def generate_html_from_regions(all_res_list, save_folder, img_name):
    """
    从layout结果生成HTML，使用绝对定位保留原始布局
    all_res_list: 包含 {'img': img, 'res': res, 'page_index': page_index} 的列表
    """
    html_content = '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {
            margin: 0;
            padding: 20px;
            font-family: 'Times New Roman', 'SimSun', serif;
            background: white;
        }
        .page {
            position: relative;
            width: 8.5in;
            height: 11in;
            margin: 20px auto;
            padding: 0.5in;
            background: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            page-break-after: always;
        }
        .region {
            position: absolute;
            box-sizing: border-box;
        }
        .region-header {
            text-align: center;
            font-weight: bold;
            font-size: 18pt;
        }
        .region-title {
            font-weight: bold;
            font-size: 16pt;
            margin-bottom: 10px;
        }
        .region-text {
            font-size: 11pt;
            line-height: 1.5;
            text-indent: 0.5in;
        }
        .region-footer {
            font-size: 10pt;
            text-align: center;
        }
        .region-table {
            font-size: 10pt;
        }
        .region-figure {
            text-align: center;
        }
        .region-figure img {
            max-width: 90%;
            max-height: 90%;
        }
        table {
            border-collapse: collapse;
            width: 100%;
        }
        table, th, td {
            border: 1px solid #000;
        }
        th, td {
            padding: 5px;
            text-align: left;
        }
    </style>
</head>
<body>
'''
    
    # 遍历每一页
    for page_idx, page_data in enumerate(all_res_list):
        res = page_data['res']
        page_index = page_data['page_index']
        
        html_content += f'    <div class="page" id="page-{page_index}">\n'
        html_content += f'        <div style="font-size: 8pt; color: #999; margin-bottom: 10px;">页 {page_index + 1}</div>\n'
        
        # 按y坐标排序
        sorted_res = sorted(res, key=lambda r: (r['bbox'][1], r['bbox'][0]))
        
        # 遍历每个region
        for region in sorted_res:
            if len(region['res']) == 0:
                continue
            
            bbox = region['bbox']
            # bbox格式: [x1, y1, x2, y2]
            # 转换为HTML坐标（像素到英寸）
            x = bbox[0] / 72  # 72 DPI
            y = bbox[1] / 72
            width = (bbox[2] - bbox[0]) / 72
            height = (bbox[3] - bbox[1]) / 72
            
            region_type = region['type'].lower()
            
            # 开始region div
            html_content += f'        <div class="region region-{region_type}" style="left: {x}in; top: {y}in; width: {width}in; height: {height}in;">\n'
            
            if region_type == 'header':
                for line in region['res']:
                    html_content += f'            <div class="region-header">{_escape_html(line["text"])}</div>\n'
            
            elif region_type == 'title':
                for line in region['res']:
                    html_content += f'            <div class="region-title">{_escape_html(line["text"])}</div>\n'
            
            elif region_type == 'footer':
                for line in region['res']:
                    html_content += f'            <div class="region-footer">{_escape_html(line["text"])}</div>\n'
            
            elif region_type == 'figure':
                # 图片处理
                img_idx = region['img_idx']
                img_path = region.get('img_path', '')
                if img_path:
                    html_content += f'            <div class="region-figure"><img src="{img_path}" alt="figure"></div>\n'
            
            elif region_type == 'table':
                if isinstance(region['res'], dict) and 'html' in region['res']:
                    html_table = region['res']['html']
                    html_content += f'            <div class="region-table">{html_table}</div>\n'
            
            else:  # 默认文本
                for line_idx, line in enumerate(region['res']):
                    text = line['text'] if isinstance(line, dict) else line
                    html_content += f'            <div class="region-text">{_escape_html(text)}</div>\n'
            
            html_content += '        </div>\n'
        
        html_content += '    </div>\n'
    
    html_content += '''</body>
</html>
'''
    return html_content


def convert_html_to_docx(html_path, docx_path):
    """
    使用pandoc将HTML转换为DOCX
    """
    try:
        # 检查pandoc是否安装
        subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
        
        # 使用pandoc转换
        cmd = [
            'pandoc',
            html_path,
            '-o', docx_path,
            '--from=html',
            '--to=docx',
            '-V', 'margin-left=0.5in',
            '-V', 'margin-right=0.5in',
            '-V', 'margin-top=0.5in',
            '-V', 'margin-bottom=0.5in'
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            logger.info(f'Successfully converted HTML to DOCX using pandoc: {docx_path}')
            return True
        else:
            logger.warning(f'Pandoc conversion failed: {result.stderr}')
            return False
    
    except FileNotFoundError:
        logger.warning('Pandoc not found, will use fallback method')
        return False
    except Exception as e:
        logger.warning(f'Error during pandoc conversion: {e}')
        return False


def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)

    flag = 1
    for i, region in enumerate(res):
        if len(region['res']) == 0:
            continue
        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        if region['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(region['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region['type'].lower() == 'title':
            doc.add_heading(region['res'][0]['text'])
        elif region['type'].lower() == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            try:
                parser.handle_table(region['res']['html'], doc)
                print('debug:::',region['res']['html'])

            except Exception as e:
                import traceback
                logger.warning(f"Failed to parse table, using text fallback: {e}\n{traceback.format_exc()}")
                # 表格处理失败时，改为插入文本内容
                paragraph = doc.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                if isinstance(region['res'], dict) and 'html' in region['res']:
                    # 尝试从 HTML 中提取文本
                    html_content = region['res']['html']
                    # 简单的 HTML 标签移除
                    import re
                    text_content = re.sub(r'<[^>]+>', '', html_content)
                    paragraph.add_run(text_content)
                else:
                    # 如果是其他格式，插入提示信息
                    paragraph.add_run("[表格识别失败，请查看生成的图片]")
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)
    logger.info('docx save to {}'.format(docx_path))


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res


def convert_info_docx_multi_page(all_res_list, save_folder, img_name, keep_reference=False):
    """
    将多页的识别结果合并到一个 Word 文档中，最大程度还原原PDF格式
    通过文本框坐标判断同一行的文本，保持原有行结构
    
    all_res_list: 包含 {'img': img, 'res': res, 'page_index': page_index} 的列表
    keep_reference: 是否保留签批区（True时单独保留reference区域图片+标记）
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)
    
    for page_idx, page_data in enumerate(all_res_list):
        img = page_data['img']
        res = page_data['res']
        page_index = page_data['page_index']
        
        # 每一页创建新的节（section），这样每页可以有独立的页眉页脚
        if page_idx > 0:
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            # 设置新节的页眉页脚不链接到上一节
            new_section.different_first_page_header_footer = False
            new_section.header.is_linked_to_previous = False
            new_section.footer.is_linked_to_previous = False
        
        # 当前页的section（第一页用默认section，后续页用新创建的section）
        current_section = doc.sections[-1]
        
        # 每页独立的页眉页脚追踪
        header_paragraph = None
        footer_paragraph = None
        
        # 按照垂直位置排序区域（从上到下）
        sorted_res = sorted(res, key=lambda r: (r['bbox'][1], r['bbox'][0]))
        
        # 处理该页的所有内容
        for i, region in enumerate(sorted_res):
            if len(region['res']) == 0:
                continue
            img_idx = region['img_idx']
            region_type = region['type'].lower()

            if region_type == 'figure':
                excel_save_folder = os.path.join(save_folder, img_name)
                img_path = os.path.join(excel_save_folder,
                                        '{}_{}.jpg'.format(region['bbox'], img_idx))
                paragraph_pic = doc.add_paragraph()
                paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph_pic.add_run("")
                try:
                    run.add_picture(img_path, width=shared.Inches(5))
                except Exception as e:
                    logger.warning(f"Failed to add picture: {e}")
                    
            elif region_type == 'header':
                # header处理：添加到当前页的页眉
                lines = group_text_by_lines(region['res'])
                
                if header_paragraph is None:
                    # 当前页第一个header，创建页眉段落
                    header = current_section.header
                    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 合并多行为横向拼接
                header_text_parts = []
                for line_items in lines:
                    line_text = merge_line_texts(line_items)
                    header_text_parts.append(line_text)
                
                merged_header_text = '  '.join(header_text_parts)
                
                # 追加到同一段落，用空格分隔
                if header_paragraph.runs:
                    header_paragraph.add_run('  ')  # 段落间分隔
                text_run = header_paragraph.add_run(merged_header_text)
                text_run.font.size = shared.Pt(10)
                        
            elif region_type == 'footer':
                # footer处理：添加到当前页的页脚
                lines = group_text_by_lines(region['res'])
                
                if footer_paragraph is None:
                    # 当前页第一个footer，创建页脚段落
                    footer = current_section.footer
                    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 合并多行为横向拼接
                footer_text_parts = []
                for line_items in lines:
                    line_text = merge_line_texts(line_items)
                    footer_text_parts.append(line_text)
                
                merged_footer_text = '  '.join(footer_text_parts)
                
                # 追加到同一段落，用空格分隔
                if footer_paragraph.runs:
                    footer_paragraph.add_run('  ')  # 段落间分隔
                text_run = footer_paragraph.add_run(merged_footer_text)
                text_run.font.size = shared.Pt(10)
                        
            elif region_type == 'title':
                # title可能有多行，需要分组处理
                lines = group_text_by_lines(region['res'])
                if lines:
                    # 创建一个标题段落
                    title_paragraph = doc.add_heading()
                    
                    # 根据region的x坐标判断对齐方式
                    # 如果左上角在页面左侧30%以内，则左对齐；否则居中
                    bbox = region['bbox']
                    page_width = bbox[2] - bbox[0] if len(bbox) >= 2 else 1000  # 大致估算页面宽度
                    
                    # 使用region的x1坐标（左上角x）
                    x1 = bbox[0]
                    # 假设页面总宽度大约是600（这是典型PDF宽度的相对值）
                    # 如果x1在左侧30%内，则左对齐
                    if x1 < 600 * 0.3:  # 左侧30%
                        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for line_idx, line_items in enumerate(lines):
                        line_text = merge_line_texts(line_items)
                        title_paragraph.add_run(line_text)
                        
                        # 除了最后一行，添加换行符
                        if line_idx < len(lines) - 1:
                            title_paragraph.add_run('\n')
                    
            elif region_type == 'table':
                parser = HtmlToDocx()
                parser.table_style = 'TableGrid'
                try:
                    parser.handle_table(region['res']['html'], doc)
                except Exception as e:
                    import traceback
                    logger.warning(f"Failed to parse table on page {page_index + 1}, using text fallback: {e}\n{traceback.format_exc()}")
                    paragraph = doc.add_paragraph()
                    if isinstance(region['res'], dict) and 'html' in region['res']:
                        html_content = region['res']['html']
                        text_content = re.sub(r'<[^>]+>', '', html_content)
                        paragraph.add_run(text_content)
                    else:
                        paragraph.add_run("[表格识别失败]")
                        
            elif region_type == 'reference':
                # Reference区域处理：根据keep_reference选项决定处理方式
                
                if keep_reference:
                    # 保留签批区模式：插入标记 + 原始图片 + 文字结果
                    
                    # 1. 插入红底黄字标记
                    flag_paragraph = doc.add_paragraph()
                    flag_run = flag_paragraph.add_run('检测到签批区')
                    flag_run.font.color.rgb = RGBColor(255, 255, 0)  # 黄字
                    flag_run.font.bold = True
                    flag_run.font.size = shared.Pt(12)
                    
                    # 设置背景色为红色 (直接使用highlight，不需要shading)
                    flag_run.font.highlight_color = 6  # WD_COLOR_INDEX.RED
                    
                    # 2. 插入reference区域的原始图片（如果存在）
                    excel_save_folder = os.path.join(save_folder, img_name)
                    img_path = os.path.join(excel_save_folder,
                                            '{}_{}.jpg'.format(region['bbox'], img_idx))
                    
                    # 检查图片文件是否存在
                    if os.path.exists(img_path):
                        paragraph_pic = doc.add_paragraph()
                        paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph_pic.add_run("")
                        try:
                            run.add_picture(img_path, width=shared.Inches(2))
                        except Exception as e:
                            logger.warning(f"Failed to add reference picture: {e}")
                    else:
                        # 图片不存在，只记录日志，不影响流程
                        logger.info(f"Reference region image not found (will only show text): {img_path}")
                    
                    # 3. 插入reference区域的文字识别结果
                    lines = group_text_by_lines(region['res'])
                    for line_idx, line_items in enumerate(lines):
                        paragraph = doc.add_paragraph()
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_after = shared.Pt(0)
                        paragraph_format.space_before = shared.Pt(0)
                        paragraph_format.line_spacing = 1.0
                        
                        line_text = merge_line_texts(line_items)
                        text_run = paragraph.add_run(line_text)
                        text_run.font.size = shared.Pt(10)
                else:
                    # 不保留签批区模式：只插入文字识别结果，不插入标记和图片
                    lines = group_text_by_lines(region['res'])
                    for line_idx, line_items in enumerate(lines):
                        paragraph = doc.add_paragraph()
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_after = shared.Pt(0)
                        paragraph_format.space_before = shared.Pt(0)
                        paragraph_format.line_spacing = 1.0
                        
                        line_text = merge_line_texts(line_items)
                        text_run = paragraph.add_run(line_text)
                        text_run.font.size = shared.Pt(10)
                    
            else:
                # 普通文本区域：按行分组还原
                lines = group_text_by_lines(region['res'])
                
                for line_idx, line_items in enumerate(lines):
                    # 每一行创建一个新段落
                    paragraph = doc.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = shared.Pt(0)  # 减少段落间距
                    paragraph_format.space_before = shared.Pt(0)
                    paragraph_format.line_spacing = 1.0  # 单倍行距
                    
                    # 合并同一行的文本
                    line_text = merge_line_texts(line_items)
                    text_run = paragraph.add_run(line_text)
                    text_run.font.size = shared.Pt(10)

    # save to docx with retry mechanism for locked files
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            # Check if file exists and is locked
            if os.path.exists(docx_path):
                try:
                    # Try to open the file in append mode to check if it's locked
                    with open(docx_path, 'ab') as f:
                        pass
                except (IOError, OSError) as e:
                    # File is locked, create a backup and use a different name
                    logger.warning('File {} is locked by another process, will save as backup'.format(docx_path))
                    timestamp = int(time.time() * 1000) % 100000
                    docx_path_backup = os.path.join(save_folder, '{}_{}_ocr.docx'.format(img_name, timestamp))
                    docx_path = docx_path_backup
                    logger.info('Using backup path: {}'.format(docx_path))
            
            doc.save(docx_path)
            logger.info('docx save to {}'.format(docx_path))
            break  # Success, exit retry loop  
        

             
        except (IOError, OSError, PermissionError) as e:
            if attempt < max_retries - 1:
                logger.warning('Failed to save docx (attempt {}/{}): {}, retrying in 1 second...'.format(
                    attempt + 1, max_retries, str(e)))
                time.sleep(1)
            else:
                # Create backup with timestamp on final failure
                timestamp = int(time.time() * 1000) % 100000
                docx_path_final = os.path.join(save_folder, '{}_{}_ocr.docx'.format(img_name, timestamp))
                try:
                    doc.save(docx_path_final)
                    logger.info('Failed to save to original path, saved to backup: {}'.format(docx_path_final))
                except Exception as final_error:
                    logger.error('Failed to save docx even after retries: {}'.format(str(final_error)))
                    raise
