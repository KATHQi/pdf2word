"""从 PDF 转 Word 最终 DOCX 中读取文本并提取法务字段。

本模块只处理已经生成的 Word 文档，不包含 OCR、PDF 渲染、图片识别、
PaddleOCR 或 PP-Structure 初始化逻辑。
"""

from __future__ import annotations

import hashlib
import json
import os
import re
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Stable domain model
# ---------------------------------------------------------------------------


@dataclass
class TextLine:
    text: str
    page: Optional[int] = None
    paragraph: Optional[int] = None
    bbox: Optional[List[float]] = None
    confidence: Optional[float] = None
    block_type: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        return {k: v for k, v in data.items() if v is not None}


@dataclass
class SourceDocument:
    filename: str
    extension: str
    text: str
    lines: List[TextLine]
    extraction_method: str
    page_count: Optional[int] = None
    sha256: Optional[str] = None
    warnings: List[str] = field(default_factory=list)

    def evidence_for(self, value: str, max_items: int = 3) -> List[Dict[str, Any]]:
        needle = normalize_for_match(value)
        if not needle:
            return []
        result: List[Dict[str, Any]] = []
        for line in self.lines:
            hay = normalize_for_match(line.text)
            if needle in hay or hay in needle:
                result.append(line.to_dict())
                if len(result) >= max_items:
                    break
        if not result:
            # For long sections, use the first line containing a stable prefix.
            prefix = needle[: min(16, len(needle))]
            for line in self.lines:
                if prefix and prefix in normalize_for_match(line.text):
                    result.append(line.to_dict())
                    if len(result) >= max_items:
                        break
        return result


@dataclass
class FieldResult:
    key: str
    frontend_label: Optional[str]
    raw_value: Any
    normalized_value: Any
    confidence: float
    mapping_level: str  # direct / conditional / unavailable / system
    evidence: List[Dict[str, Any]] = field(default_factory=list)
    note: Optional[str] = None
    value_type: str = "string"

    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        return {k: v for k, v in data.items() if v is not None}


@dataclass
class ClassificationResult:
    document_type: str
    document_type_name: str
    page_code: str
    page_name: str
    stage: str
    stage_name: str
    confidence: float
    reasons: List[str]

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# ---------------------------------------------------------------------------
# Constants and frontend profiles
# ---------------------------------------------------------------------------


DOC_TYPES: Dict[str, Dict[str, str]] = {
    "civil_complaint": {"name": "民事起诉状", "page_code": "001", "page_name": "民事起诉状"},
    "summons": {"name": "传票", "page_code": "002", "page_name": "传票"},
    "evidence_notice": {"name": "举证通知书", "page_code": "003", "page_name": "举证通知书"},
    "jurisdiction_objection": {"name": "管辖权异议申请书", "page_code": "004", "page_name": "管辖权异议"},
    "preservation_application": {"name": "财产保全申请书", "page_code": "005", "page_name": "证据 - 财产保全"},
    "defense": {"name": "答辩状", "page_code": "006", "page_name": "答辩状"},
    "appeal": {"name": "民事上诉状", "page_code": "007", "page_name": "民事上诉状"},
    "judgment": {"name": "民事判决书", "page_code": "008", "page_name": "案件判决"},
    # 程序性裁定不得写入案件判决结果；先送到原始附件，同时保留文种。
    "procedural_ruling": {"name": "程序性民事裁定书", "page_code": "010", "page_name": "原始附件"},
    "enforcement_application": {"name": "强制执行申请书", "page_code": "009", "page_name": "执行"},
    "unknown": {"name": "未识别文书", "page_code": "010", "page_name": "原始附件"},
}

# 前端审级字段只允许以下四种稳定取值。
# ``unknown`` 仅作为内部兜底状态，不应回填到前端“阶段”字段。
TRIAL_STAGE_NAMES = {
    "first_instance": "一审",
    "second_instance": "二审",
    "retrial_first_instance": "再一审",
    "retrial_second_instance": "再二审",
}

STAGE_NAMES = {
    **TRIAL_STAGE_NAMES,
    "unknown": "未知",
}

# Canonical key -> exact frontend label from the comparison report.
FRONTEND_FIELD_MAP: Dict[str, Dict[str, str]] = {
    "civil_complaint": {
        "court": "受理法院全称",
        "case_type": "案件类型",
        "labor_arbitration_no": "劳动仲裁编号",
        "case_no": "案件编号",
        "cause": "案由",
        "claim_amount": "诉讼标的额",
        "main_request_type": "主请求类型",
        "requests": "请求明细（金额/行为/标的物）",
        "facts": "法律事实",
        "laws": "法律依据",
        "contract_no": "涉案合同编号",
        "parties": "原告/被告/第三人姓名或名称",
    },
    "summons": {
        "court": "受理法院",
        "case_type": "案件类型",
        "cause": "案由",
        "case_no": "案件编号",
        "stage": "阶段",
        "summoned_person": "被传唤人",
        "summons_reason": "传唤事由",
        "issue_date": "签发日期",
        "chamber": "承办庭室",
        "hearing_time": "出庭时间",
        "hearing_place": "出庭地点",
        "judge": "法官",
        "judge_phone": "法官联系方式",
        "clerk": "书记员",
        "notes": "注意事项",
        "parties": "人员信息：当事人角色",
    },
    "evidence_notice": {
        "court": "受理法院",
        "case_type": "案件类型",
        "cause": "案由",
        "case_no": "案件编号",
        "stage": "阶段",
        "relative_deadline": "截止日期",
        "notice_target": "人员信息：通知对象",
        "evidence_submission_form": "证据提交形式",
        "recommended_evidence": "证据名称/内容/编号表",
    },
    "jurisdiction_objection": {
        "applicant": "异议提出方",
        "case_no": "案件编号",
        "case_type": "案件类型",
        "cause": "案由",
        "court": "受理法院",
        "objected_court": "被异议法院",
        "application_date": "异议提交日期",
        "jurisdiction_error_type": "管辖错误类型",
        "objection_type": "异议类型",
        "parties": "人员信息：申请人/被申请人",
    },
    "preservation_application": {
        "court": "受理法院",
        "case_type": "案件类型",
        "cause": "案由",
        "case_no": "案件编号",
        "contract_no": "涉案合同编号",
        "applicant": "申请人",
        "respondent": "被申请人",
        "preservation_target": "保全标的",
        "preservation_reason": "保全事由",
        "preservation_evidence": "保全证据",
        "guarantee_type": "担保方式",
    },
    "defense": {
        "court": "受理法院",
        "case_type": "案件类型",
        "cause": "案由",
        "case_no": "案件编号",
        "claim_amount": "诉讼标的额",
        "stage": "阶段",
        "parties": "人员信息：答辩人/被答辩人",
        "defense_opinion": "答辩意见",
    },
    "appeal": {
        "original_role": "我方为",
        "court": "受理法院全称（二审）",
        "case_type": "案件类型",
        "cause": "案由",
        "case_no": "案件编号（二审案号）",
        "claim_amount": "诉讼标的额",
        "main_request_type": "主请求类型",
        "requests": "请求明细",
        "facts": "法律事实",
        "laws": "法律依据",
        "parties": "人员信息：上诉人/被上诉人/第三人",
    },
    "judgment": {
        "case_no": "案件编号",
        "case_type": "案件类型",
        "cause": "案由",
        "contract_no": "涉案合同编号",
        "court": "受理法院",
        "claim_amount": "诉讼标的额",
        "litigation_fee": "诉讼费",
        "requests": "原告诉讼请求",
        "parties": "人员信息：原告/被告名称",
        "relative_appeal_deadline": "上诉截止日期",
        "relative_fulfillment_deadline": "判决执行截止日期",
        "judgment_result": "判决结果",
        "request_support": "支持诉讼请求",
        "case_analysis": "案件分析/案例分析",
    },
    "enforcement_application": {
        "applicant": "申请执行人",
        "respondent": "被执行人",
        "case_type": "案件类型",
        "cause": "案由",
        "basis_case_no": "执行依据文号",
        "original_case_no": "原审案号",
        "application_date": "申请执行日",
        "requested_measures": "执行措施",
        "request_details": "执行措施详情",
        "unfulfilled_principal": "未履行本金",
        "unfulfilled_interest": "未履行利息",
        "delay_interest": "迟延金",
        "parties": "人员信息：姓名/名称",
    },
    "procedural_ruling": {
        "case_no": "案件编号",
        "case_type": "案件类型",
        "court": "受理法院",
        "parties": "人员信息",
        "ruling_result": "裁定结果",
    },
}


# ---------------------------------------------------------------------------
# visual=2 handoff profiles
# ---------------------------------------------------------------------------

# The field order below is copied from “前端字段与 OCR 识别文书对照分析报告”.
# It is used only to render the visual=2 response and does not participate in
# document classification, extraction, normalization, or confidence decisions.
# The judgment page keeps the current neutral field names (for example “受理法院”
# rather than “一审受理法院”) while retaining the report-defined position.
HANDOFF_FIELD_ORDER: Dict[str, List[str]] = {'civil_complaint': ['受理法院全称',
                     '案件类型',
                     '劳动仲裁编号',
                     '案件编号',
                     '案由',
                     '诉讼标的额',
                     '一审诉讼费',
                     '诉讼费纳税人',
                     '诉讼费纳税单号',
                     '主请求类型',
                     '送达日期',
                     '案件归属业务部门',
                     '涉案合同编号',
                     '请求明细（金额/行为/标的物）',
                     '法律事实',
                     '核心内容摘要',
                     '法律依据',
                     '备注',
                     '原告/被告/第三人姓名或名称',
                     '联系电话',
                     '民族',
                     '住所地',
                     '证件类型',
                     '社会统一信用代码/身份证',
                     '委托代理人',
                     '委托人联系方式',
                     '办案人员'],
 'summons': ['受理法院',
             '案件类型',
             '案由',
             '案件编号',
             '阶段',
             '案件归属业务部门',
             '涉案合同编号',
             '被传唤人',
             '传唤事由',
             '签发日期',
             '承办庭室',
             '接收日期',
             '送达方式及签收记录',
             '人员信息：当事人角色',
             '人员信息：住所地',
             '人员信息：电话/民族/社会统一信用代码/身份证/代理人',
             '出庭时间',
             '出庭地点',
             '法官',
             '法官联系方式',
             '书记员',
             '书记员联系方式',
             '注意事项'],
 'evidence_notice': ['受理法院',
                     '案件类型',
                     '案由',
                     '案件编号',
                     '案件归属业务部门',
                     '涉案合同编号',
                     '阶段',
                     '生效日期',
                     '截止日期',
                     '送达日期',
                     '送达方式及签收记录',
                     '法官/法官联系方式',
                     '书记员/书记员联系方式',
                     '人员信息：通知对象',
                     '人员信息：原告/被告名称',
                     '人员信息：电话/民族/住所/证件/代理人',
                     '证据提交形式',
                     '证据名称/内容/编号表',
                     '备注'],
 'jurisdiction_objection': ['异议提出方',
                            '案件编号',
                            '案件类型',
                            '案由',
                            '受理法院',
                            '案件归属业务部门',
                            '涉案合同编号',
                            '被异议法院',
                            '异议提交日期',
                            '送达回证日期',
                            '管辖错误类型',
                            '异议类型',
                            '约定管辖条款',
                            '人员信息：申请人/被申请人',
                            '人员信息：联系电话',
                            '人员信息：民族',
                            '人员信息：住所地',
                            '人员信息：证件类型',
                            '人员信息：社会统一信用代码/身份证',
                            '人员信息：法定代表人',
                            '被告住所地证明',
                            '备注'],
 'preservation_application': ['受理法院',
                              '案件类型',
                              '案由',
                              '案件编号',
                              '案件归属业务部门',
                              '涉案合同编号',
                              '申请人',
                              '被申请人',
                              '保全标的',
                              '保全事由',
                              '保全证据',
                              '担保方式',
                              '人员信息：姓名/电话/民族/住所/证件',
                              '委托代理人及联系方式',
                              '备注'],
 'defense': ['受理法院',
             '案件类型',
             '案由',
             '案件编号',
             '诉讼标的额',
             '涉案合同编号',
             '答辩日期',
             '案件归属业务部门',
             '阶段',
             '人员信息：答辩人/被答辩人',
             '人员信息：电话/民族/住所/证件/代理人',
             '答辩意见',
             '备注'],
 'appeal': ['我方为',
            '受理法院全称（二审）',
            '案件类型',
            '案由',
            '案件编号（二审案号）',
            '诉讼标的额',
            '二审诉讼费',
            '诉讼费缴纳人',
            '诉讼费缴纳单号',
            '一审案件编号',
            '一审判决日期',
            '上诉材料提交日期',
            '人员信息：上诉人/被上诉人/第三人',
            '人员信息：电话/民族/住所/证件/代理人',
            '主请求类型',
            '请求明细',
            '涉案合同编号',
            '送达日期',
            '法律事实',
            '法律依据',
            '核心内容摘要',
            '判决书扫描件',
            '民事上诉状原件扫描件',
            '上诉费缴纳凭证',
            '其他证据材料',
            '案件归属业务部门',
            '备注'],
 'judgment': ['案件编号',
              '案件类型',
              '案由',
              '涉案合同编号',
              '开庭日期',
              '举证截止日期',
              '受理法院',
              '开庭地点',
              '诉讼标的额',
              '诉讼费',
              '诉讼费缴纳单号',
              '诉讼费缴纳人',
              '提交起诉材料时间',
              '立案日期',
              '原告诉讼请求',
              '人员信息：原告/被告名称',
              '人员信息：住所地',
              '人员信息：证件类型/号码',
              '人员信息：民族',
              '诉讼代理人及联系方式',
              '判决签收日期',
              '上诉截止日期',
              '判决生效日期',
              '判决执行截止日期',
              '判决执行状态',
              '判决结果',
              '支持诉讼请求',
              '案件分析/案例分析',
              '案件归属业务部门',
              '备注'],
 'enforcement_application': ['执行状态',
                             '申请执行人',
                             '被执行人',
                             '第三人',
                             '案件类型',
                             '案由',
                             '执行依据文号',
                             '执行案号',
                             '原审案号',
                             '生效日期',
                             '申请执行日',
                             '财产线索类型',
                             '财产线索详情',
                             '执行措施',
                             '执行措施详情',
                             '财产处置',
                             '人员信息：姓名/名称',
                             '人员信息：住所地',
                             '人员信息：民族',
                             '人员信息：社会统一信用代码/身份证/联系电话',
                             '人员信息：法定代表人',
                             '执行回款金额',
                             '执行回款日期',
                             '未履行本金',
                             '未履行利息',
                             '迟延金',
                             '执行和解状态',
                             '终审判决书/裁定书扫描件',
                             '执行申请书原件扫描件',
                             '执行受理通知书扫描件',
                             '财产查控相关凭证',
                             '执行和解协议/终结裁定/其他佐证',
                             '备注'],
 'procedural_ruling': ['案件编号', '案件类型', '受理法院', '人员信息：当事人', '裁定结果'],
 'unknown': ['受理法院', '案件类型', '案件编号', '案由', '涉案合同编号', '人员信息：当事人']}

# Report label -> existing canonical extraction key. Missing labels deliberately
# remain empty in visual=2; the formatter never invents values.
HANDOFF_FIELD_KEYS: Dict[str, Dict[str, str]] = {'civil_complaint': {'受理法院全称': 'court',
                     '案件类型': 'case_type',
                     '劳动仲裁编号': 'labor_arbitration_no',
                     '案件编号': 'case_no',
                     '案由': 'cause',
                     '诉讼标的额': 'claim_amount',
                     '主请求类型': 'main_request_type',
                     '涉案合同编号': 'contract_no',
                     '请求明细（金额/行为/标的物）': 'requests',
                     '法律事实': 'facts',
                     '法律依据': 'laws'},
 'summons': {'受理法院': 'court',
             '案件类型': 'case_type',
             '案由': 'cause',
             '案件编号': 'case_no',
             '阶段': 'stage',
             '涉案合同编号': 'contract_no',
             '被传唤人': 'summoned_person',
             '传唤事由': 'summons_reason',
             '签发日期': 'issue_date',
             '承办庭室': 'chamber',
             '出庭时间': 'hearing_time',
             '出庭地点': 'hearing_place',
             '法官': 'judge',
             '法官联系方式': 'judge_phone',
             '书记员': 'clerk',
             '注意事项': 'notes'},
 'evidence_notice': {'受理法院': 'court',
                     '案件类型': 'case_type',
                     '案由': 'cause',
                     '案件编号': 'case_no',
                     '涉案合同编号': 'contract_no',
                     '阶段': 'stage',
                     '截止日期': 'relative_deadline',
                     '人员信息：通知对象': 'notice_target',
                     '证据提交形式': 'evidence_submission_form',
                     '证据名称/内容/编号表': 'recommended_evidence'},
 'jurisdiction_objection': {'异议提出方': 'applicant',
                            '案件编号': 'case_no',
                            '案件类型': 'case_type',
                            '案由': 'cause',
                            '受理法院': 'court',
                            '涉案合同编号': 'contract_no',
                            '被异议法院': 'objected_court',
                            '异议提交日期': 'application_date',
                            '管辖错误类型': 'jurisdiction_error_type',
                            '异议类型': 'objection_type'},
 'preservation_application': {'受理法院': 'court',
                              '案件类型': 'case_type',
                              '案由': 'cause',
                              '案件编号': 'case_no',
                              '涉案合同编号': 'contract_no',
                              '申请人': 'applicant',
                              '被申请人': 'respondent',
                              '保全标的': 'preservation_target',
                              '保全事由': 'preservation_reason',
                              '保全证据': 'preservation_evidence',
                              '担保方式': 'guarantee_type'},
 'defense': {'受理法院': 'court',
             '案件类型': 'case_type',
             '案由': 'cause',
             '案件编号': 'case_no',
             '诉讼标的额': 'claim_amount',
             '涉案合同编号': 'contract_no',
             '阶段': 'stage',
             '答辩意见': 'defense_opinion'},
 'appeal': {'我方为': 'original_role',
            '受理法院全称（二审）': 'court',
            '案件类型': 'case_type',
            '案由': 'cause',
            '案件编号（二审案号）': 'case_no',
            '诉讼标的额': 'claim_amount',
            '主请求类型': 'main_request_type',
            '请求明细': 'requests',
            '涉案合同编号': 'contract_no',
            '法律事实': 'facts',
            '法律依据': 'laws'},
 'judgment': {'案件编号': 'case_no',
              '案件类型': 'case_type',
              '案由': 'cause',
              '涉案合同编号': 'contract_no',
              '受理法院': 'court',
              '诉讼标的额': 'claim_amount',
              '诉讼费': 'litigation_fee',
              '原告诉讼请求': 'requests',
              '上诉截止日期': 'relative_appeal_deadline',
              '判决执行截止日期': 'relative_fulfillment_deadline',
              '判决结果': 'judgment_result',
              '支持诉讼请求': 'request_support',
              '案件分析/案例分析': 'case_analysis'},
 'enforcement_application': {'申请执行人': 'applicant',
                             '被执行人': 'respondent',
                             '案件类型': 'case_type',
                             '案由': 'cause',
                             '执行依据文号': 'basis_case_no',
                             '原审案号': 'original_case_no',
                             '申请执行日': 'application_date',
                             '执行措施': 'requested_measures',
                             '执行措施详情': 'request_details',
                             '未履行本金': 'unfulfilled_principal',
                             '未履行利息': 'unfulfilled_interest',
                             '迟延金': 'delay_interest'},
 'procedural_ruling': {'案件编号': 'case_no', '案件类型': 'case_type', '受理法院': 'court', '裁定结果': 'ruling_result'},
 'unknown': {'受理法院': 'court', '案件类型': 'case_type', '案件编号': 'case_no', '案由': 'cause', '涉案合同编号': 'contract_no'}}

# Personnel fields are one-to-many. They are emitted once, at the position of
# the first personnel-related report field, while preserving the original list
# structure and translating its keys into frontend-facing Chinese labels.
HANDOFF_PERSONNEL_LABELS: Dict[str, List[str]] = {'civil_complaint': ['原告/被告/第三人姓名或名称', '联系电话', '民族', '住所地', '证件类型', '社会统一信用代码/身份证', '委托代理人', '委托人联系方式'],
 'summons': ['人员信息：当事人角色', '人员信息：住所地', '人员信息：电话/民族/社会统一信用代码/身份证/代理人'],
 'evidence_notice': ['人员信息：通知对象', '人员信息：原告/被告名称', '人员信息：电话/民族/住所/证件/代理人'],
 'jurisdiction_objection': ['人员信息：申请人/被申请人',
                            '人员信息：联系电话',
                            '人员信息：民族',
                            '人员信息：住所地',
                            '人员信息：证件类型',
                            '人员信息：社会统一信用代码/身份证',
                            '人员信息：法定代表人'],
 'preservation_application': ['人员信息：姓名/电话/民族/住所/证件', '委托代理人及联系方式'],
 'defense': ['人员信息：答辩人/被答辩人', '人员信息：电话/民族/住所/证件/代理人'],
 'appeal': ['人员信息：上诉人/被上诉人/第三人', '人员信息：电话/民族/住所/证件/代理人'],
 'judgment': ['人员信息：原告/被告名称', '人员信息：住所地', '人员信息：证件类型/号码', '人员信息：民族', '诉讼代理人及联系方式'],
 'enforcement_application': ['人员信息：姓名/名称', '人员信息：住所地', '人员信息：民族', '人员信息：社会统一信用代码/身份证/联系电话', '人员信息：法定代表人'],
 'procedural_ruling': ['人员信息：当事人'],
 'unknown': ['人员信息：当事人']}

HANDOFF_EXTRA_FIELD_LABELS: Dict[str, str] = {
    "court": "受理法院",
    "case_type": "案件类型",
    "case_no": "案件编号",
    "cause": "案由",
    "contract_no": "涉案合同编号",
    "stage": "阶段",
    "parties": "人员信息",
    "applicant": "申请人/申请执行人",
    "respondent": "被申请人/被执行人",
    "document_date": "文书签署日期",
    "summoned_address": "被传唤人住所",
    "late_consequences": "逾期举证后果",
    "issue_date": "签发日期",
    "requested_transfer_court": "拟移送法院",
    "objection_request": "异议请求",
    "objection_facts": "异议事实与理由",
    "preservation_request": "保全请求",
    "preservation_amount": "保全金额/保全限额",
    "guarantee_detail": "担保详情",
    "defense_request": "答辩请求",
    "defense_evidence": "举证材料",
    "first_instance_court": "一审法院",
    "facts_found": "法院认定事实",
    "court_reasoning": "本院认为",
    "judgment_date": "判决日期",
    "ruling_result": "裁定结果",
    "ruling_date": "裁定日期",
    "execution_facts": "执行事实与理由",
    "requested_amount": "申请执行金额",
}

# Fields that must never be auto-populated from the current document body.
BLOCKED_FIELDS: Dict[str, List[Tuple[str, str]]] = {
    "civil_complaint": [
        ("案件编号", "起诉状通常形成于立案前，不能用平台单号代替法院案号。"),
        ("一审诉讼费", "诉讼费用承担请求不等于法院核定金额。"),
        ("诉讼费纳税人", "当前文书没有实际缴费信息。"),
        ("诉讼费纳税单号", "应来自电子票据或缴费凭证。"),
        ("送达日期", "应来自送达回证或流程日志。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "summons": [
        ("接收日期", "传票签发日期不等于接收日期。"),
        ("送达方式及签收记录", "正文提示签收不代表已经实际签收。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "evidence_notice": [
        ("生效日期", "举证通知书落款日期是签发日期，不是生效日期。"),
        ("送达日期", "应来自送达回证。"),
        ("送达方式及签收记录", "应来自送达材料。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "jurisdiction_objection": [
        ("送达回证日期", "申请书不载送达回证日期。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "preservation_application": [
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "defense": [
        ("答辩日期", "文件修改时间不能代替答辩日期。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "appeal": [
        ("案件编号（二审案号）", "上诉状提交时通常尚无二审案号，不能用一审案号代替。"),
        ("二审诉讼费", "费用承担请求不等于实际缴费金额。"),
        ("诉讼费缴纳人", "请求对方承担不等于实际缴纳人。"),
        ("诉讼费缴纳单号", "应来自缴费凭证。"),
        ("送达日期", "应来自送达回证。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "judgment": [
        ("判决签收日期", "判决落款日期不等于当事人签收日期。"),
        ("判决生效日期", "需结合送达和上诉情况，不能由落款日期推断。"),
        ("诉讼费缴纳人", "费用最终负担人不等于实际缴费人。"),
        ("诉讼费缴纳单号", "应来自票据。"),
        ("判决执行状态", "属于后续执行流程字段。"),
        ("案件归属业务部门", "企业内部管理字段，OCR不得覆盖。"),
    ],
    "enforcement_application": [
        ("执行状态", "执行申请书只能表明已申请，不能表示法院执行进度。"),
        ("执行案号", "申请提交前通常尚无执行案号。"),
        ("生效日期", "文本多只表述已经生效，没有实际日期。"),
        ("执行回款金额", "属于后续执行结果。"),
        ("执行回款日期", "属于后续执行结果。"),
        ("执行和解状态", "属于后续执行流程。"),
    ],
    "procedural_ruling": [
        ("判决结果", "程序性裁定不得写入实体判决结果。"),
        ("支持诉讼请求", "程序性裁定不评价实体诉讼请求。"),
    ],
}

CAUSE_DICTIONARY = [
    "民间借贷纠纷",
    "买卖合同纠纷",
    "房屋租赁合同纠纷",
    "劳务合同纠纷",
    "劳务报酬纠纷",
    "物业服务合同纠纷",
    "网络购物合同纠纷",
    "服务合同纠纷",
    "建设工程施工合同纠纷",
    "不当得利纠纷",
    "劳动争议纠纷",
    "机动车交通事故责任纠纷",
    "交通事故纠纷",
    "承揽合同纠纷",
]

SECTION_HEADINGS = {
    "requests": ["诉讼请求", "上诉请求", "申请事项", "申请请求", "答辩请求", "保全请求", "保全标的"],
    "facts": ["事实与理由", "事实和理由", "上诉事实与理由", "上诉理由", "答辩意见", "答辩要点", "实体答辩意见", "保全事由"],
    "evidence": ["证据清单", "举证材料", "保全证据"],
    "guarantee": ["担保方式"],
    "judgment_facts": ["经审理查明", "本院经审理查明", "经审查查明"],
    "court_reasoning": ["本院认为"],
    "judgment_result": ["判决如下", "裁定如下"],
}

ROLE_LABELS = [
    "原告",
    "被告",
    "第三人",
    "申请人",
    "被申请人",
    "上诉人",
    "被上诉人",
    "答辩人",
    "被答辩人",
    "申请执行人",
    "被执行人",
    "被传唤人",
]



# ---------------------------------------------------------------------------
# Final DOCX reader
# ---------------------------------------------------------------------------


def _iter_docx_blocks(parent: _Document) -> Iterator[Union[Paragraph, Table]]:
    """按 Word 正文中的真实顺序遍历段落和表格。"""
    parent_element = parent.element.body
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _unique_table_cells(row: Any) -> List[str]:
    """去除合并单元格在 python-docx ``row.cells`` 中产生的重复引用。"""
    result: List[str] = []
    seen: set[int] = set()
    for cell in row.cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        text = normalize_text(cell.text)
        if text:
            result.append(text)
    return result


def read_docx_source(
    docx_path: Union[str, Path],
    filename: Optional[str] = None,
) -> SourceDocument:
    """读取 PDF 转 Word 界面已经导出的最终 DOCX。

    读取内容包括：
    - 正文段落；
    - 表格单元格；
    - 当恢复结果使用文本框等特殊 XML 结构且常规段落为空时，
      使用 ``w:t`` 文本节点作为兜底。

    本函数不会执行任何 OCR，也不会读取原始 PDF 或图片。
    """
    path = Path(docx_path).resolve()
    if path.suffix.lower() != ".docx":
        raise ValueError("read_docx_source 仅支持 .docx 文件")
    if not path.is_file():
        raise FileNotFoundError(f"Word 文件不存在：{path}")

    document = Document(str(path))
    lines: List[TextLine] = []
    text_parts: List[str] = []
    block_index = 0

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            if text:
                lines.append(
                    TextLine(
                        text=text,
                        paragraph=block_index,
                        block_type="paragraph",
                    )
                )
                text_parts.append(text)
                block_index += 1
            continue

        for row in block.rows:
            cells = _unique_table_cells(row)
            if not cells:
                continue
            row_text = " | ".join(cells)
            lines.append(
                TextLine(
                    text=row_text,
                    paragraph=block_index,
                    block_type="table_row",
                )
            )
            text_parts.append(row_text)
            block_index += 1

    # 部分版面恢复结果可能将文字放在文本框中，python-docx 的 paragraphs
    # 不一定能读取。只有当正文文本不足时才启用 XML 文本节点兜底，避免重复。
    primary_text = "\n".join(text_parts).strip()
    warnings: List[str] = []
    if len(normalize_for_match(primary_text)) < 30:
        xml_texts = [
            normalize_text(node.text)
            for node in document.element.xpath(".//w:t")
            if getattr(node, "text", None) and normalize_text(node.text)
        ]
        xml_text = "\n".join(xml_texts).strip()
        if len(normalize_for_match(xml_text)) > len(normalize_for_match(primary_text)):
            lines = [
                TextLine(text=value, paragraph=index, block_type="xml_text")
                for index, value in enumerate(xml_texts)
            ]
            text_parts = xml_texts
            primary_text = xml_text
            warnings.append("常规段落文本较少，已从 Word XML 文本节点读取恢复内容。")

    # 补充页眉、页脚中的文本；按节去重。
    header_footer_seen: set[str] = set()
    for section in document.sections:
        for part, block_type in (
            (section.header, "header"),
            (section.footer, "footer"),
        ):
            for paragraph in part.paragraphs:
                value = normalize_text(paragraph.text)
                key = normalize_for_match(value)
                if not value or not key or key in header_footer_seen:
                    continue
                header_footer_seen.add(key)
                lines.append(
                    TextLine(
                        text=value,
                        paragraph=block_index,
                        block_type=block_type,
                    )
                )
                text_parts.append(value)
                block_index += 1

    full_text = normalize_text("\n".join(text_parts))
    if not full_text:
        warnings.append("最终 Word 中未读取到可提取文本；请确认转换结果不是纯图片。")

    return SourceDocument(
        filename=filename or path.name,
        extension=".docx",
        text=full_text,
        lines=lines,
        extraction_method="final_docx",
        page_count=None,
        sha256=file_sha256(str(path)),
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# General normalization helpers
# ---------------------------------------------------------------------------


def normalize_for_match(value: Any) -> str:
    text = str(value or "").lower()
    text = re.sub(r"[\s\u3000]+", "", text)
    text = re.sub(r"[：:，,。；;（）()\[\]〖〗<>《》‘’“”\"'\-_/|]", "", text)
    return text


def normalize_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("（", "(").replace("）", ")")
    text = text.replace("：", ":").replace("；", ";")
    text = re.sub(r"[\u00a0\u3000]", " ", text)
    # OCR often inserts spaces between every Chinese character. Remove only such spaces.
    text = re.sub(r"(?<=[\u4e00-\u9fff])[ \t]+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", normalize_text(text))


def file_sha256(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def clean_value(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    value = re.sub(r"^[\s:：,，;；]+|[\s:：,，;；]+$", "", value)
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip() or None


def first_match(patterns: Sequence[str], text: str, flags: int = 0, group: int = 1) -> Optional[str]:
    for pattern in patterns:
        match = re.search(pattern, text, flags)
        if match:
            selected_group = group if match.lastindex and group <= match.lastindex else 0
            return clean_value(match.group(selected_group))
    return None


def all_matches(pattern: str, text: str, flags: int = 0, group: int = 1) -> List[str]:
    result: List[str] = []
    for match in re.finditer(pattern, text, flags):
        selected_group = group if match.lastindex and group <= match.lastindex else 0
        value = clean_value(match.group(selected_group))
        if value and value not in result:
            result.append(value)
    return result


def unique_preserve_order(values: Iterable[Any]) -> List[Any]:
    result: List[Any] = []
    seen = set()
    for value in values:
        marker = json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (dict, list)) else str(value)
        if marker not in seen:
            seen.add(marker)
            result.append(value)
    return result


# ---------------------------------------------------------------------------
# Chinese date and number handling
# ---------------------------------------------------------------------------


CN_DIGITS = {"〇": 0, "零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}


def chinese_year_to_int(value: str) -> Optional[int]:
    digits: List[str] = []
    for ch in value:
        if ch in CN_DIGITS:
            digits.append(str(CN_DIGITS[ch]))
        elif ch.isdigit():
            digits.append(ch)
    if len(digits) == 4:
        return int("".join(digits))
    return None


def chinese_small_number(value: str) -> Optional[int]:
    value = value.strip()
    if value.isdigit():
        return int(value)
    if not value:
        return None
    if value == "十":
        return 10
    if "十" in value:
        left, right = value.split("十", 1)
        tens = CN_DIGITS.get(left, 1) if left else 1
        ones = CN_DIGITS.get(right, 0) if right else 0
        return tens * 10 + ones
    if all(ch in CN_DIGITS for ch in value):
        return int("".join(str(CN_DIGITS[ch]) for ch in value))
    return None


def normalize_date(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    value = value.strip()
    m = re.search(r"(\d{4})\s*[年./-]\s*(\d{1,2})\s*[月./-]\s*(\d{1,2})\s*日?", value)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3))).strftime("%Y-%m-%d")
        except ValueError:
            return None
    m = re.search(r"([〇零一二三四五六七八九\d]{4})年([零一二三四五六七八九十\d]{1,3})月([零一二三四五六七八九十\d]{1,3})日", value)
    if m:
        year = chinese_year_to_int(m.group(1))
        month = chinese_small_number(m.group(2))
        day = chinese_small_number(m.group(3))
        if year and month and day:
            try:
                return datetime(year, month, day).strftime("%Y-%m-%d")
            except ValueError:
                return None
    return None


def normalize_datetime(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    date = normalize_date(value)
    if not date:
        return value.strip()
    m = re.search(r"(?:(上午|下午)\s*)?(\d{1,2})\s*[时点:]\s*(\d{1,2})?\s*分?", value)
    if not m:
        return date
    hour = int(m.group(2))
    minute = int(m.group(3) or 0)
    if m.group(1) == "下午" and hour < 12:
        hour += 12
    if m.group(1) == "上午" and hour == 12:
        hour = 0
    try:
        return f"{date}T{hour:02d}:{minute:02d}:00"
    except ValueError:
        return value.strip()


def parse_relative_deadline(text: str) -> Optional[Dict[str, Any]]:
    patterns = [
        r"(?:收到|签收|送达)(?:本通知书|本通知|文书)?之日起\s*([一二三四五六七八九十百\d]+)\s*日",
        r"(?:收到|签收|送达)(?:本通知书|本通知|文书)?后\s*([一二三四五六七八九十百\d]+)\s*日",
        r"判决生效之日起\s*([一二三四五六七八九十百\d]+)\s*日",
        r"送达之日起\s*([一二三四五六七八九十百\d]+)\s*日",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if not m:
            continue
        raw_days = m.group(1)
        days = chinese_small_number(raw_days) if not raw_days.isdigit() else int(raw_days)
        if "生效" in m.group(0):
            trigger = "judgment_effective_date"
        elif "送达" in m.group(0):
            trigger = "service_date"
        elif "签收" in m.group(0):
            trigger = "receipt_date"
        else:
            trigger = "receipt_date"
        return {
            "raw": m.group(0),
            "trigger": trigger,
            "days": days,
            "absolute_date": None,
        }
    return None


def parse_money_values(text: str) -> List[Dict[str, Any]]:
    pattern = re.compile(
        r"(?:(人民币)\s*)?(?:￥|¥)?\s*([0-9]{1,3}(?:[,，][0-9]{3})*(?:\.[0-9]+)?|[0-9]+(?:\.[0-9]+)?)\s*元"
    )
    result: List[Dict[str, Any]] = []
    for m in pattern.finditer(text):
        amount = float(m.group(2).replace(",", "").replace("，", ""))
        result.append({
            "raw": m.group(0),
            "amount": amount,
            "currency": "CNY",
            "start": m.start(),
            "end": m.end(),
        })
    return result


def request_amount_summary(request_text: Optional[str]) -> Optional[Dict[str, Any]]:
    if not request_text:
        return None
    values = parse_money_values(request_text)
    if not values:
        return None
    # Do not blindly sum when the request contains ongoing interest or alternative claims.
    uncertain = bool(re.search(r"至(?:实际)?(?:清偿|付清|履行).*日|或查封|或冻结|同等价值|暂计", request_text))
    total = sum(item["amount"] for item in values)
    return {
        "amounts": [{k: v for k, v in item.items() if k not in {"start", "end"}} for item in values],
        "sum": total,
        "currency": "CNY",
        "requires_review": uncertain,
    }


# ---------------------------------------------------------------------------
# Section, case number, cause and party parsing
# ---------------------------------------------------------------------------


def extract_case_numbers(text: str) -> List[str]:
    text = normalize_text(text)
    patterns = [
        # Longer tokens must appear first so OCR text such as“民再初/民再终”不会被短 token 截断。
        r"[（(]\s*\d{4}\s*[）)]\s*[\u4e00-\u9fffA-Za-z0-9\s]{1,28}?\s*(?:民再初|民再终|民辖终|民初|民终|民再|民申|民监|民抗|执恢|执异|执|刑初|刑终|行初|行终)\s*\d+\s*号",
        r"[（(]\s*\d{4}\s*[）)]\s*[\u4e00-\u9fffA-Za-z0-9\s]{1,40}?\s*\d+\s*号",
    ]
    result: List[str] = []
    for pattern in patterns:
        for raw in re.findall(pattern, text):
            value = re.sub(r"\s+", " ", raw).strip()
            value = value.replace("(", "（").replace(")", "）")
            value = re.sub(r"（\s*", "（", value)
            value = re.sub(r"\s*）", "）", value)
            value = re.sub(r"\s+号", "号", value)
            if value not in result:
                result.append(value)
    return result


def select_primary_case_number(
    text: str,
    case_numbers: Optional[Sequence[str]] = None,
) -> Optional[str]:
    """选择当前文书的主案号，而不是正文中引用的历史案号。

    法院裁判文书通常在标题之后先列当前案号；正文后续才会引用原审、刑事
    或关联案件案号。因此按文档出现顺序取第一个案号，和 ``common()`` 中
    “案件编号”的取值规则保持一致。最重要的是：后文引用的“民终”案号
    不得覆盖首页主案号中的“民初”。
    """
    numbers = list(case_numbers) if case_numbers is not None else extract_case_numbers(text)
    return numbers[0] if numbers else None


def _compact_case_number(case_number: Optional[str]) -> str:
    return re.sub(r"\s+", "", normalize_text(case_number or ""))


def _infer_retrial_stage(text: str, case_numbers: Sequence[str]) -> Tuple[str, str]:
    """将“民再”细分为再一审或再二审。

    “民再”代字只能证明案件属于民事再审，不能单独说明其适用一审还是
    二审程序。因此在主案号确定为“民再”后，再读取程序性表述和原生效
    裁判线索；此处的正文线索只用于细分再审，不得反向覆盖“民初/民终”。
    """
    normalized = normalize_text(text)
    compact = compact_text(normalized)

    first_procedure_patterns = [
        r"(?:按照|依照|适用)第一审程序",
        r"按第一审程序(?:审理|再审)",
        r"原来(?:系|是|属于)第一审",
        r"原生效(?:判决|裁定).{0,30}(?:民初|第一审)",
    ]
    second_procedure_patterns = [
        r"(?:按照|依照|适用)第二审程序",
        r"按第二审程序(?:审理|再审)",
        r"原来(?:系|是|属于)第二审",
        r"原生效(?:判决|裁定).{0,30}(?:民终|第二审)",
        r"(?:本院|上级人民法院).{0,20}提审",
        r"提审本案",
    ]

    if any(re.search(pattern, compact) for pattern in second_procedure_patterns):
        return "retrial_second_instance", "主案号含“民再”，且正文明确适用第二审程序或属于提审"
    if any(re.search(pattern, compact) for pattern in first_procedure_patterns):
        return "retrial_first_instance", "主案号含“民再”，且正文明确适用第一审程序"

    # 再审裁判常在开头交代原生效裁判。仅在主案号已确定为“民再”时，
    # 才允许参考后续案号；若存在“民终”，原生效裁判通常属于二审。
    referenced = [_compact_case_number(number) for number in case_numbers[1:]]
    if any("民终" in number or "民辖终" in number for number in referenced):
        return "retrial_second_instance", "主案号含“民再”，且后续原审裁判链包含“民终”案号"
    if any("民初" in number for number in referenced):
        return "retrial_first_instance", "主案号含“民再”，且后续原审裁判链仅识别到“民初”案号"

    # 前端只接受四种审级。民再案号无法进一步细分且没有程序线索时，
    # 采用“再一审”作为保守默认，并在 reasons 中明确标记，便于人工复核。
    return "retrial_first_instance", "主案号仅能确定为“民再”；未发现二审程序或提审线索，默认按“再一审”返回"


def infer_stage(
    text: str,
    case_numbers: Optional[List[str]] = None,
    document_type: Optional[str] = None,
) -> Tuple[str, List[str]]:
    """推断审级，主案号具有最高且排他的优先级。

    判定顺序：
    1. 当前文书第一个案号（主案号）；
    2. 仅当主案号无法判断时，使用标题附近的程序角色/程序名称；
    3. 对无案号的起诉状、上诉状等申请类文书，默认返回一审。

    这样可以避免一审判决书因正文引用历史二审案号、出现“上诉人”或
    “二审判决”等词而被错误识别为二审。
    """
    reasons: List[str] = []
    normalized = normalize_text(text)
    numbers = case_numbers if case_numbers is not None else extract_case_numbers(normalized)
    primary = select_primary_case_number(normalized, numbers)
    primary_compact = _compact_case_number(primary)

    if primary:
        if re.search(r"(?:执恢|执异|执)\d+号", primary_compact):
            reasons.append(f"主案号“{primary}”属于执行案号，不回填民事审级")
            return "unknown", reasons
        if "民再初" in primary_compact:
            reasons.append(f"主案号“{primary}”含“民再初”")
            return "retrial_first_instance", reasons
        if "民再终" in primary_compact:
            reasons.append(f"主案号“{primary}”含“民再终”")
            return "retrial_second_instance", reasons
        if "民初" in primary_compact:
            reasons.append(f"主案号“{primary}”含“民初”")
            return "first_instance", reasons
        if "民终" in primary_compact or "民辖终" in primary_compact:
            reasons.append(f"主案号“{primary}”含“民终/民辖终”")
            return "second_instance", reasons
        if "民再" in primary_compact:
            stage, reason = _infer_retrial_stage(normalized, numbers)
            reasons.append(reason)
            return stage, reasons
        if any(token in primary_compact for token in ["民申", "民监", "民抗"]):
            reasons.append(f"主案号“{primary}”属于再审审查/监督程序，不直接等同于再审审理审级")
            return "unknown", reasons

    # 主案号不存在或不含稳定审级代字时，只观察标题附近，避免正文中的
    # 历史案件描述污染结果。
    header = normalized[: min(len(normalized), 1800)]
    if document_type in {"civil_complaint", "jurisdiction_objection", "preservation_application", "defense", "appeal"}:
        reasons.append("未识别到可判定审级的主案号；申请/诉状类文书默认按一审返回")
        return "first_instance", reasons
    if "再审" in header:
        if re.search(r"(?:按照|依照|适用)第二审程序|提审", compact_text(header)):
            reasons.append("标题附近出现再审及第二审程序/提审表述")
            return "retrial_second_instance", reasons
        reasons.append("标题附近出现再审表述，未发现二审程序线索，默认按再一审返回")
        return "retrial_first_instance", reasons
    if "二审" in header or "上诉人" in header or "被上诉人" in header:
        reasons.append("主案号无法判断，标题附近出现二审程序角色")
        return "second_instance", reasons
    if document_type in {"summons", "evidence_notice", "judgment", "procedural_ruling"}:
        reasons.append("主案号无法判断，法院民事文书默认按一审返回")
        return "first_instance", reasons
    return "unknown", reasons


def extract_courts(text: str) -> List[str]:
    pattern = r"([\u4e00-\u9fff]{2,30}(?:高级|中级)?人民法院)"
    return unique_preserve_order(all_matches(pattern, text))


def choose_court(text: str, doc_type: str) -> Optional[str]:
    courts = extract_courts(text)
    if not courts:
        return None
    # For application/pleading documents, the court after “此致” is the filing/receiving court.
    m = re.search(r"此致\s*\n?\s*([\u4e00-\u9fff]{2,30}(?:高级|中级)?人民法院)", text)
    if m:
        return m.group(1)
    if doc_type in {"summons", "evidence_notice", "judgment", "procedural_ruling"}:
        return courts[0]
    return courts[-1]


def extract_cause(text: str) -> Optional[str]:
    for cause in CAUSE_DICTIONARY:
        if cause in text:
            return cause
    candidates = all_matches(r"([\u4e00-\u9fff]{2,24}纠纷)", text)
    if candidates:
        # Avoid overly broad snippets generated by a greedy OCR line.
        candidates.sort(key=len)
        return candidates[0]
    return None


def find_heading(text: str, headings: Sequence[str], start: int = 0) -> Optional[Tuple[int, int, str]]:
    best: Optional[Tuple[int, int, str]] = None
    for heading in headings:
        m = re.search(rf"(?m)^\s*(?:[一二三四五六七八九十0-9]+[、.．]\s*)?{re.escape(heading)}\s*(?:[:：])?\s*$", text[start:])
        if not m:
            # OCR often joins the heading with the following sentence.
            m = re.search(rf"{re.escape(heading)}\s*(?:[:：])?", text[start:])
        if m:
            item = (start + m.start(), start + m.end(), heading)
            if best is None or item[0] < best[0]:
                best = item
    return best


def extract_section(text: str, headings: Sequence[str], stop_headings: Optional[Sequence[str]] = None) -> Optional[str]:
    text = normalize_text(text)
    start_info = find_heading(text, headings)
    if not start_info:
        return None
    start = start_info[1]
    stops = list(stop_headings or [])
    if not stops:
        all_headings = [h for group in SECTION_HEADINGS.values() for h in group]
        stops = [h for h in all_headings if h not in headings]
    end = len(text)
    for heading in stops:
        info = find_heading(text, [heading], start=start)
        if info and info[0] < end:
            end = info[0]
    # Common terminal markers.
    for marker in ["\n此致", "\n审判长", "\n审 判 长", "\n申请人:", "\n上诉人:", "\n答辩人:", "\n具状人:"]:
        idx = text.find(marker, start)
        if idx >= 0 and idx < end:
            end = idx
    value = text[start:end].strip(" \n:：")
    return value or None


def split_numbered_items(text: Optional[str]) -> List[str]:
    if not text:
        return []
    normalized = normalize_text(text)
    parts = re.split(r"(?m)(?=^\s*(?:[一二三四五六七八九十]+[、.]|\(?\d+\)?[、.．]))", normalized)
    items = [re.sub(r"\s+", " ", p).strip() for p in parts if p.strip()]
    if len(items) <= 1:
        parts = re.split(r"(?=\s*(?:\d+[.、]|[一二三四五六七八九十]+、))", normalized)
        items = [re.sub(r"\s+", " ", p).strip() for p in parts if p.strip()]
    return items


_NUMBERED_ARABIC_ITEM_RE = re.compile(r"(?m)^\s*(?P<number>\d{1,2})\s*[.．、)]\s*")
_REQUEST_ACTION_WORDS = (
    "判令", "请求", "责令", "支付", "偿还", "返还", "退还", "赔偿",
    "承担", "清偿", "解除", "确认", "撤销", "停止", "腾退", "交付",
    "办理", "查封", "冻结", "扣押", "保全",
)


def _numbered_blocks(text: Optional[str]) -> Tuple[str, List[Tuple[int, str]]]:
    """将阿拉伯数字编号段落拆为 ``(前缀, [(序号, 原文), ...])``。"""
    normalized = normalize_text(text or "")
    matches = list(_NUMBERED_ARABIC_ITEM_RE.finditer(normalized))
    if not matches:
        return normalized, []

    prefix = normalized[: matches[0].start()].strip()
    items: List[Tuple[int, str]] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(normalized)
        block = normalized[match.start():end].strip()
        if block:
            items.append((int(match.group("number")), block))
    return prefix, items


def _take_first_numbered_sentence(text: str) -> Optional[Tuple[int, str, str]]:
    """从文本开头取出一个编号项，返回 ``(序号, 项文本, 剩余文本)``。

    诉讼请求通常以分号结束；为兼容最后一项，也接受句号。只截取首个
    终止标点，避免把后续“事实与理由”正文一起吞入编号项。
    """
    normalized = normalize_text(text)
    match = _NUMBERED_ARABIC_ITEM_RE.match(normalized)
    if not match:
        return None

    body_start = match.end()
    terminator = re.search(r"[;。](?=\s*(?:\n|$))", normalized[body_start:])
    if terminator:
        end = body_start + terminator.end()
    else:
        next_item = _NUMBERED_ARABIC_ITEM_RE.search(normalized, body_start)
        if not next_item:
            return None
        end = next_item.start()

    item = normalized[:end].strip()
    rest = normalized[end:].lstrip(" \n")
    return int(match.group("number")), item, rest


def _looks_like_request_item(text: str) -> bool:
    compact = compact_text(text)
    return any(word in compact for word in _REQUEST_ACTION_WORDS)


def repair_civil_complaint_sections(
    requests: Optional[str],
    facts: Optional[str],
) -> Tuple[Optional[str], Optional[str], bool]:
    """修复 PDF 对象流乱序造成的诉请/事实串段。

    典型异常是：诉讼请求实际为 1、2、3 三项，但 PDF 文字对象流顺序
    变成 ``1 -> 3 -> 事实与理由 -> 2 -> 事实正文``。此时普通标题切片
    会把第 2 项误放入“法律事实”。本函数只在“事实”开头编号项能够填补
    诉请编号缺口（或紧接现有最大编号），且内容具有请求动作语义时移动，
    从而避免误伤事实正文中正常的编号叙述。
    """
    normalized_requests = normalize_text(requests or "") or None
    normalized_facts = normalize_text(facts or "") or None
    if not normalized_requests or not normalized_facts:
        return normalized_requests, normalized_facts, False

    prefix, request_items = _numbered_blocks(normalized_requests)
    if not request_items:
        return normalized_requests, normalized_facts, False

    numbers = {number for number, _ in request_items}
    moved_items: List[Tuple[int, str]] = []
    remaining_facts = normalized_facts

    while True:
        candidate = _take_first_numbered_sentence(remaining_facts)
        if not candidate:
            break
        number, item, rest = candidate
        max_number = max(numbers) if numbers else 0
        missing_numbers = set(range(1, max_number + 1)) - numbers
        fills_sequence = number in missing_numbers or number == max_number + 1
        if not fills_sequence or not _looks_like_request_item(item):
            break
        moved_items.append((number, item))
        numbers.add(number)
        remaining_facts = rest

    if not moved_items:
        return normalized_requests, normalized_facts, False

    merged: Dict[int, str] = {}
    for number, item in request_items + moved_items:
        # 保留第一次出现的版本；视觉顺序修复只补缺，不覆盖已有诉请。
        merged.setdefault(number, item)

    ordered_items = [merged[number] for number in sorted(merged)]
    repaired_requests = "\n".join(
        [part for part in ([prefix] if prefix else []) + ordered_items if part]
    ).strip()
    repaired_facts = remaining_facts.strip() or None
    return repaired_requests or None, repaired_facts, True


def extract_laws(text: str) -> List[str]:
    laws = all_matches(r"《([^》]{2,80})》", text)
    articles = all_matches(r"(第[一二三四五六七八九十百千万0-9]+条(?:第[一二三四五六七八九十0-9]+款)?)", text)
    result = [f"《{name}》" for name in laws] + articles
    return unique_preserve_order(result)


def extract_contract_numbers(text: str) -> List[str]:
    patterns = [
        r"(?:合同编号|合同号|协议编号)\s*[:：]?\s*([A-Za-z0-9][A-Za-z0-9._\-/]{3,50})",
        r"(?:编号为|编号)\s*[:：]?\s*([A-Za-z]{1,10}[-_/][A-Za-z0-9._\-/]{3,50})",
    ]
    result: List[str] = []
    for pattern in patterns:
        result.extend(all_matches(pattern, text))
    return unique_preserve_order(result)


def extract_label_value(text: str, labels: Sequence[str], max_len: int = 120) -> Optional[str]:
    labels_pattern = "|".join(re.escape(label) for label in labels)
    patterns = [
        rf"(?:{labels_pattern})\s*[:：]\s*([^\n]{{1,{max_len}}})",
        rf"(?:{labels_pattern})[ \t]+([^\n]{{1,{max_len}}})",
        # OCR/PDF extraction may remove the visual gap between label and value.
        rf"(?:{labels_pattern})\s*[:：]?\s*([^\n]{{1,{max_len}}})",
    ]
    return first_match(patterns, text)


def parse_party_block(role: str, block: str) -> Dict[str, Any]:
    block = normalize_text(block)
    # Remove parenthetical role suffix from label while retaining it as current role detail.
    role_detail = None
    m_role = re.search(r"[（(]([^）)]{1,30})[）)]", role)
    if m_role:
        role_detail = m_role.group(1)
    base_role = re.sub(r"[（(].*?[）)]", "", role)

    name = block.split("\n", 1)[0]
    name = re.split(r"[,，。;；]", name, maxsplit=1)[0]
    name = re.sub(r"^(?:为|系)", "", name).strip()
    name = name[:80]

    address = first_match([
        r"(?:住所地|住所|住址|户籍住址|现住址)\s*[:：]?\s*(.*?)(?=[,，]\s*(?:身份证|公民身|统一社会|联系|电话|手机|法定|负责)|[。；;\n]|$)",
        r"(?:住所地|住所|住址|户籍住址|现住址)\s*[:：]?\s*([^。；;\n]{2,120})",
    ], block)
    phone = first_match([
        r"(?:联系电话|电话|手机)\s*[:：]?\s*([0-9Xx*\-—()（）]{5,30})",

    ], block)
    id_no = first_match([
        r"(?:公民身份号码|身份证号码|身份证号)\s*[:：]?\s*([0-9Xx*\\\-]{6,30})",
    ], block)
    credit_code = first_match([
        r"统一社会信用代码\s*[:：]?\s*([0-9A-ZxX*]{8,30})",
    ], block)
    gender = first_match([r"[,，]\s*([男女])\s*[,，]"], block)
    nationality = first_match([r"[,，]\s*([\u4e00-\u9fff]{1,8}族)\s*[,，]"], block)
    birth_raw = first_match([r"(\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)出生"], block)
    legal_representative = first_match([
        r"法定代表人\s*[:：]?\s*([^,，。；;\n]{2,40})",
    ], block)
    representative_title = first_match([
        r"法定代表人\s*[:：]?\s*[^,，。；;\n]{2,40}[,，]\s*([^。；;\n]{1,30})",
    ], block)
    agent = first_match([
        r"委托(?:诉讼)?代理人\s*[:：]?\s*([^,，。；;\n]{2,50})",
    ], block)

    company_markers = ["公司", "集团", "中心", "事务所", "委员会", "合伙企业", "合作社"]
    subject_type = "organization" if any(marker in name for marker in company_markers) or credit_code else "person"
    data: Dict[str, Any] = {
        "role": base_role,
        "current_role_detail": role_detail,
        "name": name,
        "subject_type": subject_type,
        "address": address,
        "phone": phone,
        "gender": gender,
        "nationality": nationality,
        "birth_date": normalize_date(birth_raw),
        "id_type": "统一社会信用代码" if credit_code else ("身份证" if id_no else None),
        "id_number": credit_code or id_no,
        "legal_representative": legal_representative,
        "representative_title": representative_title,
        "agent": agent,
        "raw": block,
    }
    return {k: v for k, v in data.items() if v is not None}


def extract_parties(text: str) -> List[Dict[str, Any]]:
    text = normalize_text(text)
    # Match role labels and capture until the next role/section marker.
    role_pattern = "|".join(sorted((re.escape(r) for r in ROLE_LABELS), key=len, reverse=True))
    pattern = re.compile(
        rf"(?m)^\s*(?P<role>{role_pattern})(?P<detail>\s*[（(][^）)]{{1,30}}[）)])?\s*[:：]\s*(?P<body>.*?)(?=^\s*(?:{role_pattern})\s*(?:[（(][^）)]{{1,30}}[）)])?\s*[:：]|^\s*(?:诉讼请求|上诉请求|申请事项|申请请求|事实与理由|事实和理由|答辩意见|答辩要点|证据清单|保全请求|保全标的|保全事由|保全证据|担保方式|本院认为|经审理查明|判决如下|裁定如下)\s*$|\Z)",
        re.S,
    )
    parties: List[Dict[str, Any]] = []
    for m in pattern.finditer(text):
        role = m.group("role") + (m.group("detail") or "")
        body = m.group("body").strip()
        if not body:
            continue
        party = parse_party_block(role, body)
        if party.get("name"):
            parties.append(party)

    # A signature line such as “申请人：周凯” may repeat the richer party block.
    # Keep one record per role/name and prefer the record carrying more attributes.
    merged: Dict[Tuple[str, str], Dict[str, Any]] = {}
    order: List[Tuple[str, str]] = []
    for party in parties:
        marker = (str(party.get("role") or ""), str(party.get("name") or ""))
        if marker not in merged:
            merged[marker] = party
            order.append(marker)
        elif len(party) > len(merged[marker]):
            merged[marker] = party
    return [merged[marker] for marker in order]


def extract_issue_or_signature_date(text: str) -> Optional[Tuple[str, str]]:
    # Prefer a date near the end of the document.
    candidates: List[Tuple[int, str]] = []
    patterns = [
        r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日",
        r"[〇零一二三四五六七八九]{4}年[零一二三四五六七八九十]{1,3}月[零一二三四五六七八九十]{1,3}日",
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            candidates.append((m.start(), m.group(0)))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0])
    raw = candidates[-1][1]
    return raw, normalize_date(raw) or raw


def classify_request_type(request_text: Optional[str]) -> List[str]:
    if not request_text:
        return []
    mapping = [
        ("给付", ["支付", "偿还", "返还", "赔偿", "退还", "承担", "清偿"]),
        ("撤销", ["撤销"]),
        ("改判", ["改判"]),
        ("发回重审", ["发回重审"]),
        ("解除", ["解除合同", "解除"]),
        ("确认", ["确认"]),
        ("行为", ["腾退", "交付", "停止", "更换", "办理退货"]),
        ("保全", ["冻结", "查封", "扣押", "保全"]),
    ]
    result = [name for name, keywords in mapping if any(keyword in request_text for keyword in keywords)]
    return unique_preserve_order(result)


def classify_jurisdiction(text: str) -> Tuple[Optional[str], Optional[str]]:
    """将管辖权异议归入前端允许的三个固定枚举。

    判断优先级为：专属管辖错误 > 级别管辖错误 > 地域管辖错误。
    优先级不能颠倒，因为专属管辖理由中也经常出现“所在地”等地域词。
    没有足够依据时返回 ``(None, None)``，避免强行填充错误类别。
    """
    normalized_text = re.sub(r"\s+", "", text or "")
    if not normalized_text:
        return None, None

    # 先识别文书中直接出现的标准枚举或异议类型。
    explicit_rules = [
        ("专属管辖错误", "专属管辖异议", ["专属管辖错误", "专属管辖异议"]),
        ("级别管辖错误", "级别管辖异议", ["级别管辖错误", "级别管辖异议"]),
        ("地域管辖错误", "地域管辖异议", ["地域管辖错误", "地域管辖异议"]),
    ]
    for error_type, objection_type, keywords in explicit_rules:
        if any(keyword in normalized_text for keyword in keywords):
            return error_type, objection_type

    # 专属管辖具有排他性，必须优先于一般地域管辖判断。
    exclusive_keywords = [
        "专属管辖",
        "不动产所在地人民法院管辖",
        "不动产所在地法院管辖",
        "不动产纠纷",
        "港口所在地人民法院管辖",
        "港口作业纠纷",
        "遗产所在地人民法院管辖",
        "主要遗产所在地人民法院管辖",
        "继承遗产纠纷",
    ]
    if any(keyword in normalized_text for keyword in exclusive_keywords):
        return "专属管辖错误", "专属管辖异议"

    # 级别管辖需要出现明确的法院层级、受理范围或标的额依据；
    # 仅出现“基层/中级/高级人民法院”的法院名称，不足以单独认定。
    level_keywords = [
        "级别管辖",
        "管辖级别",
        "越级管辖",
        "超越级别管辖",
        "不符合级别管辖",
        "不属于基层人民法院管辖",
        "不应由基层人民法院管辖",
        "应由中级人民法院管辖",
        "应由高级人民法院管辖",
        "应由最高人民法院管辖",
        "超过基层人民法院受理范围",
        "超过中级人民法院受理范围",
        "诉讼标的额超过",
        "标的额超过",
    ]
    level_patterns = [
        r"(?:本案|该案).{0,20}(?:应当|应由|依法应由)(?:中级|高级|最高)人民法院管辖",
        r"(?:基层|中级)人民法院.{0,20}(?:无权|不具有|没有)(?:级别)?管辖权",
        r"(?:诉讼)?标的额.{0,20}(?:超过|超出).{0,20}(?:受理|管辖)(?:范围|标准)",
    ]
    if any(keyword in normalized_text for keyword in level_keywords) or any(
        re.search(pattern, normalized_text) for pattern in level_patterns
    ):
        return "级别管辖错误", "级别管辖异议"

    # 其余具有明确连接点依据的管辖争议，归入地域管辖错误。
    territorial_keywords = [
        "地域管辖",
        "被告住所地",
        "原告住所地",
        "申请人住所地",
        "被申请人住所地",
        "经常居住地",
        "合同履行地",
        "合同签订地",
        "侵权行为地",
        "侵权结果发生地",
        "公司住所地",
        "票据支付地",
        "保险标的物所在地",
        "运输始发地",
        "运输目的地",
        "财产所在地",
        "约定管辖",
        "协议管辖",
    ]
    territorial_patterns = [
        r"(?:受诉|原审|该|本)法院.{0,20}(?:无权|不具有|没有)管辖权",
        r"不属于.{0,30}人民法院管辖",
        r"应由.{2,30}人民法院管辖",
        r"移送(?:至|到)?.{2,30}人民法院(?:审理|管辖)",
    ]
    if any(keyword in normalized_text for keyword in territorial_keywords) or any(
        re.search(pattern, normalized_text) for pattern in territorial_patterns
    ):
        return "地域管辖错误", "地域管辖异议"

    return None, None


def classify_guarantee_type(text: str) -> Optional[str]:
    if "保险" in text and "保函" in text:
        return "诉讼财产保全责任保险保函"
    if "现金保证金" in text or "现金担保" in text:
        return "现金担保"
    if "房产" in text and "担保" in text:
        return "实物担保-房产"
    if "担保公司" in text or "专业担保公司" in text:
        return "担保公司保函"
    if "保函" in text:
        return "保函"
    return None


# ---------------------------------------------------------------------------
# Document classification
# ---------------------------------------------------------------------------


CLASSIFICATION_RULES: Dict[str, List[Tuple[str, float]]] = {
    "civil_complaint": [("民事起诉状", 8.0), ("诉讼请求", 2.0), ("具状人", 1.0)],
    "summons": [("传票", 8.0), ("被传唤人", 3.0), ("应到时间", 2.0), ("传唤事由", 2.0)],
    "evidence_notice": [("举证通知书", 10.0), ("举证期限", 2.0), ("逾期举证", 1.0)],
    "jurisdiction_objection": [("管辖权异议申请书", 10.0), ("移送", 2.0), ("无管辖权", 2.0)],
    "preservation_application": [("财产保全申请书", 10.0), ("保全事由", 2.0), ("担保方式", 1.0)],
    "defense": [("民事二审答辩状", 10.0), ("民事答辩状", 9.0), ("答辩意见", 2.0), ("答辩请求", 2.0)],
    "appeal": [("民事上诉状", 10.0), ("上诉请求", 2.0), ("不服", 1.0), ("上诉人", 1.0)],
    "judgment": [("民事判决书", 10.0), ("判决如下", 4.0), ("本院认为", 1.0)],
    "procedural_ruling": [("民事裁定书", 8.0), ("裁定如下", 3.0), ("转为适用普通程序", 8.0)],
    "enforcement_application": [("强制执行申请书", 10.0), ("申请执行", 2.0), ("执行费用", 1.0)],
}


def classify_document(text: str, filename: str = "", hint: Optional[str] = None) -> ClassificationResult:
    normalized = normalize_text(text)
    scores: Dict[str, float] = {key: 0.0 for key in CLASSIFICATION_RULES}
    reasons: Dict[str, List[str]] = {key: [] for key in CLASSIFICATION_RULES}

    if hint and hint in DOC_TYPES and hint != "unknown":
        scores[hint] += 12.0
        reasons[hint].append("调用方显式指定 document_type")

    for doc_type, rules in CLASSIFICATION_RULES.items():
        for keyword, weight in rules:
            if keyword in normalized:
                scores[doc_type] += weight
                reasons[doc_type].append(f"正文包含“{keyword}”")
            elif keyword in filename:
                scores[doc_type] += weight * 0.35
                reasons[doc_type].append(f"文件名包含“{keyword}”（仅作低权重辅助）")

    # Critical semantic override: a conversion-to-ordinary-procedure ruling is not a judgment.
    if "转为适用普通程序" in normalized or ("民事裁定书" in normalized and "普通程序" in normalized):
        scores["procedural_ruling"] += 20.0
        scores["judgment"] = min(scores["judgment"], 1.0)
        reasons["procedural_ruling"].append("识别到程序转换裁定，禁止归入实体判决")

    best_type = max(scores, key=scores.get)
    best_score = scores[best_type]
    sorted_scores = sorted(scores.values(), reverse=True)
    margin = best_score - (sorted_scores[1] if len(sorted_scores) > 1 else 0.0)
    if best_score < 3.0:
        best_type = "unknown"
        confidence = 0.25
        final_reasons = ["正文和文件名均未达到文种识别阈值"]
    else:
        confidence = min(0.99, 0.55 + best_score / 30.0 + max(0.0, margin) / 50.0)
        final_reasons = reasons[best_type]

    case_numbers = extract_case_numbers(normalized)
    stage, stage_reasons = infer_stage(normalized, case_numbers, document_type=best_type)
    final_reasons = unique_preserve_order(final_reasons + stage_reasons)

    profile = DOC_TYPES[best_type]
    return ClassificationResult(
        document_type=best_type,
        document_type_name=profile["name"],
        page_code=profile["page_code"],
        page_name=profile["page_name"],
        stage=stage,
        stage_name=STAGE_NAMES.get(stage, stage),
        confidence=round(confidence, 4),
        reasons=final_reasons,
    )


# ---------------------------------------------------------------------------
# Field extraction engine
# ---------------------------------------------------------------------------


class ExtractionBuilder:
    def __init__(self, source: SourceDocument, classification: ClassificationResult):
        self.source = source
        self.classification = classification
        self.fields: Dict[str, FieldResult] = {}
        self.warnings: List[str] = list(source.warnings)

    @property
    def doc_type(self) -> str:
        return self.classification.document_type

    def add(
        self,
        key: str,
        raw: Any,
        normalized: Any = None,
        confidence: float = 0.9,
        level: str = "direct",
        note: Optional[str] = None,
        value_type: str = "string",
        evidence_value: Optional[str] = None,
    ) -> None:
        if raw is None or raw == "" or raw == [] or raw == {}:
            return
        if normalized is None:
            normalized = raw
        label = FRONTEND_FIELD_MAP.get(self.doc_type, {}).get(key)
        if isinstance(evidence_value, str):
            evidence = self.source.evidence_for(evidence_value)
        elif isinstance(raw, str):
            evidence = self.source.evidence_for(raw)
        else:
            evidence = []
        self.fields[key] = FieldResult(
            key=key,
            frontend_label=label,
            raw_value=raw,
            normalized_value=normalized,
            confidence=round(max(0.0, min(1.0, confidence)), 4),
            mapping_level=level,
            evidence=evidence,
            note=note,
            value_type=value_type,
        )

    def common(self) -> None:
        text = self.source.text
        doc_type = self.doc_type
        case_numbers = extract_case_numbers(text)
        court = choose_court(text, doc_type)
        cause = extract_cause(text)
        contracts = extract_contract_numbers(text)
        parties = extract_parties(text)

        if court:
            self.add("court", court, confidence=0.95, level="direct")
        self.add("case_type", "民事", confidence=0.78, level="conditional", note="由文书标题、案号及案由推断。")
        primary_case_number = select_primary_case_number(text, case_numbers)
        if primary_case_number:
            # “案件编号”和审级必须使用同一个主案号，禁止被正文引用案号覆盖。
            self.add("case_no", primary_case_number, confidence=0.96, level="direct")
        if cause:
            self.add("cause", cause, confidence=0.94, level="direct")
        if contracts:
            self.add("contract_no", contracts, confidence=0.88, level="conditional", value_type="array")
        if self.classification.stage in TRIAL_STAGE_NAMES:
            self.add(
                "stage",
                self.classification.stage_name,
                normalized=self.classification.stage,
                confidence=self.classification.confidence,
                level="conditional",
                note="审级以当前文书主案号为最高优先级；正文引用的历史案号不得覆盖主案号。",
            )
        if parties:
            self.add("parties", parties, confidence=0.89, level="direct", value_type="array")

    def build(self) -> Dict[str, Any]:
        safe: Dict[str, Any] = {}
        review: Dict[str, Any] = {}
        for field_result in self.fields.values():
            if not field_result.frontend_label:
                continue
            if field_result.mapping_level == "direct" and field_result.confidence >= 0.82:
                safe[field_result.frontend_label] = field_result.normalized_value
            else:
                review[field_result.frontend_label] = {
                    "value": field_result.normalized_value,
                    "confidence": field_result.confidence,
                    "reason": field_result.note or "条件对应字段，需业务复核。",
                }

        blocked = [
            {"frontend_label": label, "reason": reason}
            for label, reason in BLOCKED_FIELDS.get(self.doc_type, [])
        ]

        return {
            "classification": self.classification.to_dict(),
            "fields": {key: value.to_dict() for key, value in self.fields.items()},
            "frontend_patch": {
                "page_code": self.classification.page_code,
                "page_name": self.classification.page_name,
                "safe_autofill": safe,
                "review_required": review,
                "blocked": blocked,
            },
            "warnings": unique_preserve_order(self.warnings),
        }


def extract_common_date(builder: ExtractionBuilder, key: str, level: str = "direct", confidence: float = 0.9) -> None:
    date_info = extract_issue_or_signature_date(builder.source.text)
    if date_info:
        raw, normalized = date_info
        builder.add(key, raw, normalized=normalized, confidence=confidence, level=level)


def extract_civil_complaint(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    requests = extract_section(text, ["诉讼请求"], ["事实与理由", "事实和理由", "证据清单"])
    facts = extract_section(text, ["事实与理由", "事实和理由"], ["证据清单"])
    evidence = extract_section(text, ["证据清单"], [])

    requests, facts, repaired = repair_civil_complaint_sections(requests, facts)
    if repaired:
        builder.warnings.append(
            "检测到 PDF 文字对象流导致诉讼请求编号串入事实与理由，已按编号连续性和请求语义修复。"
        )

    amount = request_amount_summary(requests)
    request_types = classify_request_type(requests)
    # 甲方约定：“法律依据”直接使用“证据清单”段落内容。
    laws = evidence

    if requests:
        builder.add(
            "requests",
            raw=requests,
            normalized=requests,
            confidence=0.95,
            level="direct",
            value_type="text",
        )
    if facts:
        builder.add("facts", facts, confidence=0.94, level="direct", value_type="text")
    if amount:
        builder.add(
            "claim_amount",
            amount,
            confidence=0.72 if amount["requires_review"] else 0.86,
            level="conditional",
            note="金额由诉讼请求中的明示金额聚合；持续计息、备选请求或行为请求需人工确认。",
            value_type="money_summary",
            evidence_value=requests,
        )
    if request_types:
        builder.add(
            "main_request_type",
            request_types,
            confidence=0.78,
            level="conditional",
            value_type="array",
            evidence_value=requests,
        )
    if laws:
        builder.add(
            "laws",
            laws,
            confidence=0.95,
            level="direct",
            note="按甲方字段规则，法律依据直接取证据清单段落。",
            value_type="text",
        )
    extract_common_date(builder, "document_date", level="direct", confidence=0.9)



def extract_summons(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    summoned = extract_label_value(text, ["被传唤人"])
    reason = extract_label_value(text, ["传唤事由"])
    hearing_raw = extract_label_value(text, ["应到时间"])
    place = extract_label_value(text, ["应到处所"])
    judge = extract_label_value(text, ["承办法官", "审判员", "审判长"])
    judge_phone = extract_label_value(text, ["联系电话", "法官联系方式"])
    clerk = extract_label_value(text, ["书记员"])
    notes = extract_section(text, ["注意事项"], ["承办法官", "审判长", "审判员", "书记员"])
    address = extract_label_value(text, ["住所地", "住所"])

    if summoned:
        builder.add("summoned_person", summoned, confidence=0.97, level="direct")
    if reason:
        builder.add("summons_reason", reason, confidence=0.96, level="direct")
    if hearing_raw:
        builder.add("hearing_time", hearing_raw, normalized=normalize_datetime(hearing_raw), confidence=0.97, level="direct", value_type="datetime")
    if place:
        builder.add("hearing_place", place, confidence=0.97, level="direct")
        chamber = first_match([r"(第?[一二三四五六七八九十0-9]+(?:审判庭|法庭)|[\u4e00-\u9fff]{1,20}(?:审判庭|法庭))"], place)
        if chamber:
            builder.add("chamber", chamber, confidence=0.74, level="conditional", note="从应到处所中抽取；地点不一定等于承办庭室。")
    if judge:
        builder.add("judge", judge, confidence=0.92, level="direct")
    if judge_phone:
        builder.add("judge_phone", judge_phone, confidence=0.9, level="direct")
    if clerk:
        builder.add("clerk", clerk, confidence=0.92, level="direct")
    if notes:
        builder.add("notes", raw=notes, normalized=split_numbered_items(notes), confidence=0.93, level="direct", value_type="array")
    if address:
        builder.add("summoned_address", address, confidence=0.9, level="direct")
    extract_common_date(builder, "issue_date", level="direct", confidence=0.9)


def extract_evidence_notice(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    target = None
    # Usually the addressee is the first standalone line after the case number.
    m = re.search(r"(?:号)\s*\n\s*([^:\n]{1,60})[:：]\s*\n", text)
    if m:
        target = clean_value(m.group(1))
    if not target:
        target = first_match([r"本院受理原告[^\n]{1,40}诉(?:你方|你)"], text)
    relative = parse_relative_deadline(text)
    submission_form = first_match([
        r"(证据材料应当提交复印件[^。]{0,160}。)",
        r"(全部证据需[^。]{0,160}。)",
        r"(证据材料须[^。]{0,160}。)",
    ], text)
    recommended = first_match([
        r"(?:应当|应|需|须)提交([^。]{3,250}(?:证据|材料|记录|凭证)[^。]{0,100})",
    ], text)
    late = first_match([
        r"(逾期[^。]{2,180}(?:后果|采纳|罚款|裁判|风险)[^。]{0,80}。?)",
        r"(未能提供证据[^。]{2,180}。)",
    ], text)

    if target:
        builder.add("notice_target", target, confidence=0.9, level="direct")
    if relative:
        builder.add(
            "relative_deadline",
            relative["raw"],
            normalized=relative,
            confidence=0.86,
            level="conditional",
            note="仅保存相对期限；缺少实际送达/签收日期时不得生成绝对截止日。",
            value_type="relative_deadline",
        )
    if submission_form:
        builder.add("evidence_submission_form", submission_form, confidence=0.9, level="direct", value_type="text")
    if recommended:
        builder.add(
            "recommended_evidence",
            recommended,
            confidence=0.72,
            level="conditional",
            note="这是法院要求或建议提交的证据，不代表当事人已经上传。",
            value_type="text",
        )
    if late:
        builder.add("late_consequences", late, confidence=0.9, level="direct", value_type="text")
    # Deliberately expose issue date as a non-frontend metadata field, never map to effective date.
    extract_common_date(builder, "issue_date", level="direct", confidence=0.9)


def extract_jurisdiction_objection(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    parties = builder.fields.get("parties")
    applicant = None
    respondent = None
    if parties:
        for party in parties.normalized_value:
            if party.get("role") == "申请人" and not applicant:
                applicant = party.get("current_role_detail")
            if party.get("role") == "被申请人" and not respondent:
                respondent = party.get("current_role_detail")
    request_text = extract_section(text, ["申请事项", "申请请求"], ["事实与理由", "事实和理由"])
    requested_court = first_match([
        r"移送\s*([\u4e00-\u9fff]{2,30}(?:高级|中级)?人民法院)\s*审理",
        r"移送至\s*([\u4e00-\u9fff]{2,30}(?:高级|中级)?人民法院)",
    ], request_text or text)
    facts = extract_section(text, ["事实与理由", "事实和理由"], [])
    jurisdiction_basis_text = "\n".join(part for part in [request_text, facts] if part) or text
    error_type, objection_type = classify_jurisdiction(jurisdiction_basis_text)
    current_court = choose_court(text, "jurisdiction_objection")

    if applicant:
        builder.add("applicant", applicant, confidence=0.95, level="direct")
    if respondent:
        builder.add("respondent", respondent, confidence=0.92, level="direct")
    if current_court:
        builder.add("objected_court", current_court, confidence=0.95, level="direct")
    if requested_court:
        builder.add("requested_transfer_court", requested_court, confidence=0.96, level="direct")
    if request_text:
        builder.add("objection_request", request_text, confidence=0.95, level="direct", value_type="text")
    if facts:
        builder.add("objection_facts", facts, confidence=0.94, level="direct", value_type="text")
    if error_type:
        builder.add(
            "jurisdiction_error_type",
            error_type,
            confidence=0.82,
            level="conditional",
            note="仅在地域管辖错误、级别管辖错误、专属管辖错误三个固定枚举中归类。",
        )
    if objection_type:
        builder.add("objection_type", objection_type, confidence=0.76, level="conditional")
    extract_common_date(builder, "application_date", level="direct", confidence=0.92)


def extract_preservation_receiving_court(text: str) -> Optional[str]:
    """从财产保全样本开头的“法院：具体法院”元数据中提取受理法院。"""
    compact = compact_text(text)
    header = compact[:500]

    court_suffix = (
        r"(?:人民法院|海事法院|知识产权法院|"
        r"互联网法院|金融法院|铁路运输法院)"
    )

    return first_match(
        [
            # 标准样本：法院：杭州市余杭区人民法院
            rf"法院[:：|丨｜]?([\u4e00-\u9fff]{{2,40}}?{court_suffix})"
            rf"(?=版式|[｜|】\]]|$)",

            # OCR 未识别“版式”前的分隔符
            rf"法院[:：|丨｜]?([\u4e00-\u9fff]{{2,40}}?{court_suffix})",

            # 极端情况下，“法院”标签被漏识别，但开头仍存在完整法院名称
            rf"([\u4e00-\u9fff]{{2,40}}?{court_suffix})",
        ],
        header,
    )


def extract_preservation_application(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    parties = builder.fields.get("parties")
    applicant = None
    respondent = None
    if parties:
        for party in parties.normalized_value:
            if party.get("role") == "申请人" and not applicant:
                applicant = party.get("name")
            if party.get("role") == "被申请人" and not respondent:
                respondent = party.get("name")
    target = extract_section(text, ["保全请求", "保全标的"], ["保全事由", "事实与理由"])
    reason = extract_section(text, ["保全事由"], ["保全证据", "担保方式"])
    evidence = extract_section(text, ["保全证据"], ["担保方式"])
    guarantee = extract_section(text, ["担保方式"], [])
    guarantee_type = classify_guarantee_type(guarantee or text)
    amount = request_amount_summary(target)
    receiving_court = extract_preservation_receiving_court(text)

    # 覆盖 common() 中通用法院识别可能从“恳请人民法院/请求人民法院”
    # 误抽出的结果；这里只采用甲方样本首行元数据中的明确法院。
    if receiving_court:
        builder.add(
            "court",
            receiving_court,
            confidence=0.99,
            level="direct",
            note="从文书首行【案由｜法院｜版式】元数据中的‘法院’字段提取。",
        )

    if applicant:
        builder.add("applicant", applicant, confidence=0.9, level="conditional" if "脱敏" in applicant else "direct")
    if respondent:
        builder.add("respondent", respondent, confidence=0.9, level="conditional" if "脱敏" in respondent else "direct")
    if target:
        builder.add("preservation_target", target, confidence=0.94, level="direct", value_type="text")
        builder.add("preservation_request", target, confidence=0.94, level="direct", value_type="text")
    if amount:
        builder.add("preservation_amount", amount, confidence=0.82, level="conditional", note="前端当前缺少独立保全金额字段，建议新增。", value_type="money_summary", evidence_value=target)
    if reason:
        builder.add("preservation_reason", reason, confidence=0.94, level="direct", value_type="text")
    if evidence:
        builder.add("preservation_evidence", raw=evidence, normalized=split_numbered_items(evidence), confidence=0.93, level="direct", value_type="array")
    if guarantee_type:
        builder.add("guarantee_type", guarantee_type, confidence=0.9, level="direct", evidence_value=guarantee or guarantee_type)
    if guarantee:
        builder.add("guarantee_detail", guarantee, confidence=0.9, level="direct", note="前端当前只有担保方式，建议增加担保详情。", value_type="text")


def extract_defense(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    opinion = extract_section(text, ["实体答辩意见", "答辩意见", "答辩要点"], ["举证材料", "答辩请求"])
    request_text = extract_section(text, ["答辩请求"], [])
    evidence = extract_section(text, ["举证材料"], ["答辩请求"])
    amount = request_amount_summary(opinion)

    if opinion:
        full = opinion + (("\n答辩请求\n" + request_text) if request_text else "")
        builder.add("defense_opinion", full, confidence=0.94, level="direct", value_type="text")
    if request_text:
        builder.add("defense_request", raw=request_text, normalized=split_numbered_items(request_text), confidence=0.94, level="direct", value_type="array")
    if evidence:
        builder.add("defense_evidence", raw=evidence, normalized=split_numbered_items(evidence), confidence=0.9, level="direct", value_type="array")
    if amount:
        builder.add("claim_amount", amount, confidence=0.62, level="conditional", note="答辩状往往仅反驳部分金额，不能视为完整诉讼标的额。", value_type="money_summary")


def extract_appeal(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    requests = extract_section(text, ["上诉请求"], ["事实与理由", "上诉事实与理由", "上诉理由"])
    facts = extract_section(text, ["上诉事实与理由", "事实与理由", "上诉理由"], [])
    amount = request_amount_summary(requests)
    request_types = classify_request_type(requests)
    laws = extract_laws(text)
    original_role = first_match([
        r"上诉人\s*[（(]原审([^）)]{1,20})[）)]",
    ], text)
    first_instance_court = first_match([
        r"不服\s*([\u4e00-\u9fff]{2,30}人民法院)(?:作出|的)?(?:一审)?民事判决",
        r"不服\s*([\u4e00-\u9fff]{2,30}人民法院)[^。]{0,60}一审判决",
    ], text)

    if original_role:
        builder.add("original_role", f"原审{original_role}", confidence=0.95, level="direct")
    if first_instance_court:
        builder.add("first_instance_court", first_instance_court, confidence=0.92, level="direct", note="前端当前缺少独立一审法院字段。")
    if requests:
        builder.add("requests", raw=requests, normalized=split_numbered_items(requests), confidence=0.95, level="direct", value_type="array")
    if facts:
        builder.add("facts", facts, confidence=0.94, level="direct", value_type="text")
    if amount:
        builder.add("claim_amount", amount, confidence=0.7, level="conditional", note="上诉请求含非定额损失或行为请求时需人工确认。", value_type="money_summary")
    if request_types:
        builder.add("main_request_type", request_types, confidence=0.8, level="conditional", value_type="array")
    if laws:
        builder.add("laws", laws, confidence=0.78, level="conditional", note="仅提取明示法律，不补造依据。", value_type="array")
    # Remove case number field if only an original first-instance number is present.
    case_field = builder.fields.get("case_no")
    if case_field and "民初" in str(case_field.normalized_value) and "民终" not in str(case_field.normalized_value):
        del builder.fields["case_no"]
        builder.warnings.append("只识别到一审案号；未将其写入二审案号字段。")


def extract_litigation_fees(text: str) -> Optional[List[Dict[str, Any]]]:
    """提取判决书正文中最先出现的一项诉讼费用。

    判决书可能同时列出案件受理费、保全费、上诉费等多个结果。根据业务
    规则，无论当前案件属于一审、二审、再一审还是再二审，均只保留原文
    顺序中的第一个匹配结果。
    """
    fee_types = ["案件受理费", "受理费", "保全费", "上诉费", "执行费"]
    fee_type_pattern = "|".join(
        re.escape(fee_type) for fee_type in sorted(fee_types, key=len, reverse=True)
    )
    pattern = re.compile(
        rf"(?P<type>{fee_type_pattern})\s*(?P<amount>[0-9,，.]+)\s*元"
    )
    match = pattern.search(text)
    if not match:
        return None
    return [{
        "type": match.group("type"),
        "amount": float(match.group("amount").replace(",", "").replace("，", "")),
        "currency": "CNY",
        "raw": match.group(0),
    }]


def compare_request_support(requests: Optional[str], judgment: Optional[str]) -> Optional[Dict[str, Any]]:
    if not requests or not judgment:
        return None
    request_items = split_numbered_items(requests)
    rejected = bool(re.search(r"驳回[^。；;]*(?:其他|全部)?诉讼请求", judgment))
    supported_terms = sum(1 for term in ["支付", "赔偿", "返还", "承担", "偿还"] if term in judgment)
    if rejected and supported_terms:
        result = "部分支持"
    elif rejected and not supported_terms:
        result = "不支持/驳回"
    elif supported_terms:
        result = "支持或主要支持"
    else:
        result = "无法稳定判断"
    return {
        "result": result,
        "request_item_count": len(request_items),
        "basis": "对比原告诉讼请求与判决主文关键词；属于规则推断。",
    }


def extract_judgment(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    requests = first_match([
        r"原告[^。]{0,80}提出诉讼请求\s*[:：]?\s*(.*?)(?=事实与理由|被告[^。]{0,40}(?:答辩|未提出答辩)|经审理查明|本院认为)",
    ], text, flags=re.S)
    if not requests:
        requests = extract_section(text, ["诉讼请求"], ["事实与理由", "经审理查明", "本院认为"])
    facts_found = extract_section(text, ["经审理查明", "本院经审理查明"], ["本院认为"])
    reasoning = extract_section(text, ["本院认为"], ["判决如下", "裁定如下"])
    judgment = extract_section(text, ["判决如下"], [])
    fees = extract_litigation_fees(text)
    amount = request_amount_summary(requests)
    if amount and amount.get("amounts"):
        # 《民事判决书》的诉讼标的额若匹配到多个金额，只保留原文中第一个。
        # 该规则仅在判决书提取流程中生效，不改变起诉状、上诉状等文书的
        # request_amount_summary 聚合行为。
        first_amount = amount["amounts"][0]
        amount = {
            "amounts": [first_amount],
            "sum": first_amount["amount"],
            "currency": first_amount.get("currency", "CNY"),
            "requires_review": amount.get("requires_review", False),
            "selection_rule": "first_occurrence",
        }
    support = compare_request_support(requests, judgment)
    date_info = extract_issue_or_signature_date(text)
    appeal_deadline = parse_relative_deadline(first_match([r"(如不服本判决.*?上诉。?)"], text, flags=re.S) or "")
    fulfillment = parse_relative_deadline(judgment or text)

    if requests:
        builder.add("requests", raw=requests, normalized=split_numbered_items(requests), confidence=0.93, level="direct", value_type="array")
    if amount:
        builder.add("claim_amount", amount, confidence=0.78, level="conditional", note="根据原告诉讼请求中的明示金额计算。", value_type="money_summary")
    if facts_found:
        builder.add("facts_found", facts_found, confidence=0.92, level="direct", value_type="text")
    if reasoning:
        builder.add("court_reasoning", reasoning, confidence=0.92, level="direct", value_type="text")
        builder.add("case_analysis", reasoning[:3000], confidence=0.72, level="conditional", note="当前返回本院认为原文，不自动生成无依据的案例分析。", value_type="text")
    if judgment:
        builder.add("judgment_result", raw=judgment, normalized=split_numbered_items(judgment), confidence=0.96, level="direct", value_type="array")
    if fees:
        builder.add("litigation_fee", fees, confidence=0.94, level="direct", value_type="array")
    if support:
        builder.add("request_support", support, confidence=0.68, level="conditional", note="属于诉请与判项规则比对，不是判决书中的独立键值。", value_type="object")
    if date_info:
        raw, normalized = date_info
        builder.add("judgment_date", raw, normalized=normalized, confidence=0.92, level="direct", note="判决日期单独返回，绝不写入签收日期或生效日期。", value_type="date")
    if appeal_deadline:
        builder.add("relative_appeal_deadline", appeal_deadline["raw"], normalized=appeal_deadline, confidence=0.8, level="conditional", note="缺少实际送达日时，只保留相对规则。", value_type="relative_deadline")
    if fulfillment:
        builder.add("relative_fulfillment_deadline", fulfillment["raw"], normalized=fulfillment, confidence=0.78, level="conditional", note="缺少生效日时，不生成绝对履行截止日。", value_type="relative_deadline")


def extract_procedural_ruling(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    result = extract_section(text, ["裁定如下"], [])
    if result:
        result = re.sub(r"\n?(?:审判员|审判长|书记员)[^\n]*.*$", "", result, flags=re.S).strip()
        result = re.sub(r"\n?(?:\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日|[〇零一二三四五六七八九]{4}年[零一二三四五六七八九十]{1,3}月[零一二三四五六七八九十]{1,3}日)\s*$", "", result).strip()
        builder.add("ruling_result", raw=result, normalized=split_numbered_items(result), confidence=0.96, level="direct", value_type="array")
    date_info = extract_issue_or_signature_date(text)
    if date_info:
        raw, normalized = date_info
        builder.add("ruling_date", raw, normalized=normalized, confidence=0.92, level="direct", value_type="date")
    builder.warnings.append("该文书属于程序性裁定，已强制路由到“原始附件”，不会生成实体判决结果。")


def extract_enforcement_application(builder: ExtractionBuilder) -> None:
    text = builder.source.text
    parties = builder.fields.get("parties")
    applicant = None
    respondent = None
    if parties:
        for party in parties.normalized_value:
            if party.get("role") in {"申请人", "申请执行人"} and not applicant:
                applicant = party.get("name")
            if party.get("role") in {"被申请人", "被执行人"} and not respondent:
                respondent = party.get("name")
    request_text = extract_section(text, ["申请事项", "申请请求"], ["事实与理由", "事实和理由"])
    facts = extract_section(text, ["事实与理由", "事实和理由"], [])
    case_numbers = extract_case_numbers(text)
    basis = None
    for number in case_numbers:
        if "执" not in number:
            basis = number
            break
    amount = request_amount_summary(request_text)
    principal = first_match([
        r"(?:借款本金|本金|货款|工程款|劳务报酬|拖欠工资)\s*(?:人民币)?\s*([0-9,，.]+)\s*元",
    ], request_text or text)
    interest_text = first_match([
        r"((?:逾期|法定|合同约定)?利息[^；;。\n]{0,120})",
    ], request_text or text)
    delay_text = first_match([
        r"(迟延履行期间(?:加倍)?债务利息[^；;。\n]{0,120})",
    ], request_text or text)

    if applicant:
        builder.add("applicant", applicant, confidence=0.94, level="direct")
    if respondent:
        builder.add("respondent", respondent, confidence=0.94, level="direct")
    if basis:
        builder.add("basis_case_no", basis, confidence=0.9, level="conditional", note="仅在申请书明确引用生效裁判文号时填写。")
        builder.add("original_case_no", basis, confidence=0.86, level="conditional", note="需结合一审/二审裁判链确认该文号语义。")
    if request_text:
        builder.add("request_details", raw=request_text, normalized=split_numbered_items(request_text), confidence=0.94, level="conditional", note="这是申请采取的措施，不代表法院已经实际执行。", value_type="array")
        builder.add("requested_measures", classify_request_type(request_text), confidence=0.76, level="conditional", note="标记为申请措施，不作为已执行措施。", value_type="array")
    if facts:
        builder.add("execution_facts", facts, confidence=0.92, level="direct", value_type="text")
    if amount:
        builder.add("requested_amount", amount, confidence=0.8, level="conditional", value_type="money_summary")
    if principal:
        builder.add("unfulfilled_principal", principal, normalized={"amount": float(principal.replace(",", "").replace("，", "")), "currency": "CNY"}, confidence=0.72, level="conditional", note="仅当申请书明确表述尚未履行或未支付时可作为初始值。", value_type="money")
    if interest_text:
        builder.add("unfulfilled_interest", interest_text, confidence=0.66, level="conditional", note="保留计算规则，不在缺少利率/起止日时估算具体金额。", value_type="text")
    if delay_text:
        builder.add("delay_interest", delay_text, confidence=0.72, level="conditional", note="仅表示请求计算迟延履行利息，不填造金额。", value_type="text")
    extract_common_date(builder, "application_date", level="direct", confidence=0.92)
    # No execution case number or execution status is produced from an application.
    case_field = builder.fields.get("case_no")
    if case_field:
        del builder.fields["case_no"]


TYPE_EXTRACTORS = {
    "civil_complaint": extract_civil_complaint,
    "summons": extract_summons,
    "evidence_notice": extract_evidence_notice,
    "jurisdiction_objection": extract_jurisdiction_objection,
    "preservation_application": extract_preservation_application,
    "defense": extract_defense,
    "appeal": extract_appeal,
    "judgment": extract_judgment,
    "procedural_ruling": extract_procedural_ruling,
    "enforcement_application": extract_enforcement_application,
}


def extract_legal_document(
    source: SourceDocument,
    document_type_hint: Optional[str] = None,
) -> Dict[str, Any]:
    source.text = normalize_text(source.text)
    classification = classify_document(source.text, source.filename, hint=document_type_hint)
    builder = ExtractionBuilder(source, classification)
    builder.common()
    extractor = TYPE_EXTRACTORS.get(classification.document_type)
    if extractor:
        extractor(builder)
    else:
        builder.warnings.append("文种未识别，仅返回原始附件路由和通用元数据。")
    result = builder.build()
    result["source"] = {
        "filename": source.filename,
        "extension": source.extension,
        "sha256": source.sha256,
        "page_count": source.page_count,
        "extraction_method": source.extraction_method,
        "text_length": len(source.text),
    }
    return result




# ---------------------------------------------------------------------------
# visual=2 handoff formatter (output-only; extraction logic remains unchanged)
# ---------------------------------------------------------------------------

_PARTY_FIELD_LABELS: List[Tuple[str, str]] = [
    ("role", "当事人角色"),
    ("current_role_detail", "原审/当前角色说明"),
    ("name", "姓名或名称"),
    ("subject_type", "主体类型"),
    ("phone", "联系电话"),
    ("nationality", "民族"),
    ("address", "住所地"),
    ("gender", "性别"),
    ("birth_date", "出生日期"),
    ("id_type", "证件类型"),
    ("id_number", "社会统一信用代码/身份证"),
    ("legal_representative", "法定代表人"),
    ("representative_title", "法定代表人职务"),
    ("agent", "委托代理人"),
    ("raw", "原始文本"),
]


def _handoff_has_value(value: Any) -> bool:
    return value is not None and value != "" and value != [] and value != {}


def _handoff_conf_level(field_data: Optional[Dict[str, Any]]) -> str:
    if not field_data or not _handoff_has_value(field_data.get("normalized_value")):
        return "not-extracted"
    confidence = float(field_data.get("confidence") or 0.0)
    if field_data.get("mapping_level") == "direct" and confidence >= 0.82:
        return "safe_autofill"
    return "review_required"


def _handoff_field_payload(
    field_data: Optional[Dict[str, Any]],
    value_override: Any = None,
) -> Dict[str, Any]:
    if not field_data:
        return {"value": "", "conf_level": "not-extracted"}
    value = field_data.get("normalized_value") if value_override is None else value_override
    if not _handoff_has_value(value):
        return {"value": "", "conf_level": "not-extracted"}
    return {"value": value, "conf_level": _handoff_conf_level(field_data)}


def _translate_party_items(value: Any) -> List[Dict[str, Any]]:
    if not isinstance(value, list):
        return []
    translated: List[Dict[str, Any]] = []
    for item in value:
        if not isinstance(item, dict):
            continue
        row: Dict[str, Any] = {}
        for source_key, chinese_label in _PARTY_FIELD_LABELS:
            if source_key not in item:
                continue
            field_value = item.get(source_key)
            if source_key == "subject_type":
                field_value = {
                    "person": "自然人",
                    "organization": "法人或其他组织",
                }.get(str(field_value), field_value)
            row[chinese_label] = field_value
        if row:
            translated.append(row)
    return translated


def _synthetic_personnel(
    document_type: str,
    fields: Dict[str, Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Build a minimal list only from fields already extracted by existing logic."""
    specs: Dict[str, List[Tuple[str, str]]] = {
        "summons": [("被传唤人", "summoned_person")],
        "evidence_notice": [("通知对象", "notice_target")],
        "preservation_application": [("申请人", "applicant"), ("被申请人", "respondent")],
        "enforcement_application": [("申请执行人", "applicant"), ("被执行人", "respondent")],
    }
    rows: List[Dict[str, Any]] = []
    used: List[str] = []
    for role, key in specs.get(document_type, []):
        field_data = fields.get(key)
        if not field_data or not _handoff_has_value(field_data.get("normalized_value")):
            continue
        row: Dict[str, Any] = {
            "当事人角色": role,
            "姓名或名称": field_data.get("normalized_value"),
        }
        used.append(key)
        if document_type == "summons":
            address = fields.get("summoned_address")
            if address and _handoff_has_value(address.get("normalized_value")):
                row["住所地"] = address.get("normalized_value")
                used.append("summoned_address")
        rows.append(row)
    return rows, unique_preserve_order(used)


def _handoff_personnel_payload(
    document_type: str,
    fields: Dict[str, Dict[str, Any]],
) -> Tuple[Dict[str, Any], List[str]]:
    parties = fields.get("parties")
    if parties and _handoff_has_value(parties.get("normalized_value")):
        translated = _translate_party_items(parties.get("normalized_value"))
        if translated:
            return _handoff_field_payload(parties, translated), ["parties"]

    rows, used = _synthetic_personnel(document_type, fields)
    if not rows:
        return {"value": "", "conf_level": "not-extracted"}, []
    levels = [_handoff_conf_level(fields.get(key)) for key in used]
    conf_level = "safe_autofill" if levels and all(level == "safe_autofill" for level in levels) else "review_required"
    return {"value": rows, "conf_level": conf_level}, used


def build_visual2_content(result: Dict[str, Any]) -> Dict[str, Any]:
    """Render one extraction result in the ordered frontend handoff contract.

    This function only rearranges values already present in ``result["fields"]``.
    It never reruns extraction, changes confidence, fills blocked fields, or infers
    values that the existing extraction pipeline did not produce.
    """
    classification = result.get("classification") or {}
    document_type = str(classification.get("document_type") or "unknown")
    fields: Dict[str, Dict[str, Any]] = result.get("fields") or {}
    order = HANDOFF_FIELD_ORDER.get(document_type, HANDOFF_FIELD_ORDER["unknown"])
    key_map = HANDOFF_FIELD_KEYS.get(document_type, HANDOFF_FIELD_KEYS["unknown"])
    personnel_labels = set(HANDOFF_PERSONNEL_LABELS.get(document_type, []))

    content: Dict[str, Any] = {}
    consumed: set[str] = set()
    personnel_written = False

    for frontend_label in order:
        if frontend_label in personnel_labels:
            if not personnel_written:
                personnel_payload, personnel_keys = _handoff_personnel_payload(document_type, fields)
                content["人员信息"] = personnel_payload
                consumed.update(personnel_keys)
                for personnel_label in personnel_labels:
                    personnel_key = key_map.get(personnel_label)
                    if personnel_key and personnel_key in fields:
                        consumed.add(personnel_key)
                personnel_written = True
            continue

        field_key = key_map.get(frontend_label)
        field_data = fields.get(field_key) if field_key else None
        if field_key == "stage" and field_data:
            content[frontend_label] = _handoff_field_payload(
                field_data,
                field_data.get("raw_value"),
            )
        else:
            content[frontend_label] = _handoff_field_payload(field_data)
        if field_key and field_data:
            consumed.add(field_key)

    # ``stage`` already has a dedicated top-level key in visual=2. Avoid repeating
    # it in the supplemental area when the report page itself has no “阶段” field.
    consumed.add("stage")

    extras: Dict[str, Any] = {}
    for field_key, field_data in fields.items():
        if field_key in consumed:
            continue
        value = field_data.get("normalized_value")
        if not _handoff_has_value(value):
            continue
        chinese_label = (
            field_data.get("frontend_label")
            or HANDOFF_EXTRA_FIELD_LABELS.get(field_key)
            or field_key
        )
        # Avoid accidental collision if two technical fields share a frontend label.
        final_label = str(chinese_label)
        suffix = 2
        while final_label in extras:
            final_label = f"{chinese_label}（{suffix}）"
            suffix += 1
        if field_key == "parties":
            extras[final_label] = _handoff_field_payload(
                field_data,
                _translate_party_items(value),
            )
        else:
            extras[final_label] = _handoff_field_payload(field_data)

    if extras:
        content["识别补充字段"] = extras
    return content


# ---------------------------------------------------------------------------
# API schema introspection
# ---------------------------------------------------------------------------


def api_schema() -> Dict[str, Any]:
    return {
        "trial_stages": [
            {"code": code, "name": name}
            for code, name in TRIAL_STAGE_NAMES.items()
        ],
        "document_types": {
            key: {
                **value,
                "frontend_fields": FRONTEND_FIELD_MAP.get(key, {}),
                "blocked_fields": [
                    {"frontend_label": label, "reason": reason}
                    for label, reason in BLOCKED_FIELDS.get(key, [])
                ],
            }
            for key, value in DOC_TYPES.items()
        },
        "mapping_levels": {
            "direct": "文书存在明确同义字段，可在置信度足够时安全自动回填。",
            "conditional": "需要推断、计算、角色匹配或样本并非全部具备，必须进入复核区。",
            "unavailable": "当前文书没有该信息或语义不等价，保持空值。",
            "system": "平台或人工流程字段，Word 文书提取不得覆盖。",
        },
        "visual_2_contract": {
            "result_item_keys": ["filename", "document_type", "page_name", "stage", "index", "content"],
            "field_value": {"value": "提取值；未提取时为空字符串", "conf_level": "safe_autofill/review_required/not-extracted"},
            "order": "按照《前端字段与 OCR 识别文书对照分析报告》的页面字段顺序",
        },
        "field_contract": {
            "raw_value": "最终 Word 中提取到的原始值",
            "normalized_value": "标准化后值",
            "confidence": "0-1 置信度",
            "mapping_level": "direct/conditional/unavailable/system",
            "evidence": "页码、段落、bbox和原始文本定位",
        },
    }
