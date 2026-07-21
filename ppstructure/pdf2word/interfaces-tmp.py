"""Legal-document OCR extraction API.

Drop-in replacement for the experimental ``interfaces-tmp.py`` entry point.
It keeps OCR/layout recognition inside the existing pdf2word project, while
exposing a stable, evidence-aware JSON contract for an external frontend.

Run:
    python interfaces-tmp.py

Primary endpoints:
    GET  /api/v1/health
    GET  /api/v1/schemas
    POST /api/v1/ocr/extract       multipart/form-data, field name: files
    POST /api/v1/ocr/extract-text  application/json (development/testing)

The API deliberately separates:
- document classification and trial-stage inference;
- raw OCR values and normalized values;
- safe auto-fill fields and fields requiring human review;
- document-derived fields and system/manual fields.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import tempfile
import threading
import uuid
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import fitz  # PyMuPDF
from flask import Flask, jsonify, request
from flask_cors import CORS
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.utils import secure_filename

try:
    from .legal_ocr_core import SourceDocument, TextLine, api_schema, extract_legal_document
except ImportError:  # Direct script execution from this directory.
    from legal_ocr_core import SourceDocument, TextLine, api_schema, extract_legal_document


LOGGER = logging.getLogger("legal_ocr_api")
logging.basicConfig(
    level=os.getenv("OCR_LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
)


def env_bool(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on", "y"}


def env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except ValueError:
        return default


APP_HOST = os.getenv("OCR_API_HOST", "0.0.0.0")
APP_PORT = env_int("OCR_API_PORT", 8005)
MAX_CONTENT_MB = env_int("OCR_MAX_CONTENT_MB", 80)
MAX_FILES = env_int("OCR_MAX_FILES", 12)
PDF_TEXT_PAGE_MIN_CHARS = env_int("OCR_PDF_TEXT_PAGE_MIN_CHARS", 60)
PDF_RENDER_DPI = env_int("OCR_PDF_RENDER_DPI", 180)
API_KEY = os.getenv("OCR_API_KEY", "").strip()
ALLOWED_ORIGINS = [
    item.strip()
    for item in os.getenv("OCR_CORS_ORIGINS", "*").split(",")
    if item.strip()
]

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tif",
    ".tiff",
    ".docx",
}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_MB * 1024 * 1024
app.json.ensure_ascii = False
CORS(app, resources={r"/api/*": {"origins": ALLOWED_ORIGINS}})


# ---------------------------------------------------------------------------
# OCR engine: lazy, process-local and lock-protected
# ---------------------------------------------------------------------------


class OCREngine:
    """Lazy PP-Structure wrapper.

    Paddle inference objects are expensive and not guaranteed to be safe for
    concurrent calls. A lock prevents multiple Flask request threads from
    using the same engine at the same time.
    """

    def __init__(self) -> None:
        self._engine: Any = None
        self._load_error: Optional[str] = None
        self._lock = threading.RLock()

    @property
    def loaded(self) -> bool:
        return self._engine is not None

    @property
    def load_error(self) -> Optional[str]:
        return self._load_error

    def _build_kwargs(self) -> Dict[str, Any]:
        kwargs: Dict[str, Any] = {
            "show_log": env_bool("OCR_SHOW_LOG", False),
            "lang": os.getenv("OCR_LANG", "ch"),
            "layout": env_bool("OCR_ENABLE_LAYOUT", True),
            "table": env_bool("OCR_ENABLE_TABLE", True),
            "ocr": True,
            "use_gpu": env_bool("OCR_USE_GPU", True),
        }
        model_envs = {
            "layout_model_dir": "OCR_LAYOUT_MODEL_DIR",
            "table_model_dir": "OCR_TABLE_MODEL_DIR",
            "det_model_dir": "OCR_DET_MODEL_DIR",
            "rec_model_dir": "OCR_REC_MODEL_DIR",
            "cls_model_dir": "OCR_CLS_MODEL_DIR",
        }
        for argument, env_name in model_envs.items():
            value = os.getenv(env_name, "").strip()
            if value:
                kwargs[argument] = value
        if os.getenv("OCR_GPU_MEM"):
            kwargs["gpu_mem"] = env_int("OCR_GPU_MEM", 4000)
        return kwargs

    def ensure_loaded(self) -> Any:
        if self._engine is not None:
            return self._engine
        if self._load_error:
            raise RuntimeError(self._load_error)
        with self._lock:
            if self._engine is not None:
                return self._engine
            try:
                from paddleocr import PPStructure

                kwargs = self._build_kwargs()
                LOGGER.info("Loading PPStructure with configured models")
                self._engine = PPStructure(**kwargs)
                LOGGER.info("PPStructure loaded")
                return self._engine
            except Exception as exc:  # pragma: no cover - depends on runtime models.
                self._load_error = f"PPStructure initialization failed: {exc}"
                LOGGER.exception(self._load_error)
                raise RuntimeError(self._load_error) from exc

    def recognize(self, image: Any, page: int) -> List[TextLine]:
        engine = self.ensure_loaded()
        with self._lock:
            result = engine(image)
        return parse_ppstructure_result(result, page=page)


OCR_ENGINE = OCREngine()


# ---------------------------------------------------------------------------
# Source decoding
# ---------------------------------------------------------------------------


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_line_text(text: Any) -> str:
    return " ".join(str(text or "").replace("\u3000", " ").split()).strip()


def redact_red_seal(image: Any) -> Any:
    """Lightweight optional red-seal suppression before OCR.

    This only whitens strongly red pixels. The original file is never changed,
    and the option is off by default because red text may occasionally be
    meaningful evidence.
    """
    try:
        import cv2
        import numpy as np

        hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
        low1 = cv2.inRange(hsv, np.array([0, 70, 45]), np.array([12, 255, 255]))
        low2 = cv2.inRange(hsv, np.array([165, 70, 45]), np.array([180, 255, 255]))
        mask = cv2.bitwise_or(low1, low2)
        output = image.copy()
        output[mask > 0] = 255
        return output
    except Exception:
        LOGGER.warning("Red-seal suppression failed; continuing with original image", exc_info=True)
        return image


def _bbox_from_region(region: Dict[str, Any]) -> Optional[List[float]]:
    bbox = region.get("bbox") or region.get("box")
    if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
        try:
            return [float(item) for item in bbox]
        except (TypeError, ValueError):
            return None
    return None


def _text_from_nested(value: Any) -> List[Tuple[str, Optional[float], Optional[List[float]]]]:
    """Collect text from the heterogeneous output shapes used by PPStructure."""
    collected: List[Tuple[str, Optional[float], Optional[List[float]]]] = []
    if value is None:
        return collected
    if isinstance(value, str):
        text = normalize_line_text(value)
        if text and not text.lstrip().startswith("<table"):
            collected.append((text, None, None))
        return collected
    if isinstance(value, dict):
        direct = value.get("text") or value.get("transcription")
        if isinstance(direct, str):
            confidence = value.get("confidence", value.get("score"))
            try:
                confidence = float(confidence) if confidence is not None else None
            except (TypeError, ValueError):
                confidence = None
            bbox = value.get("text_region") or value.get("bbox") or value.get("box")
            bbox_value: Optional[List[float]] = None
            if isinstance(bbox, (list, tuple)):
                # A quadrilateral is converted to its enclosing rectangle.
                try:
                    if len(bbox) == 4 and all(isinstance(x, (int, float)) for x in bbox):
                        bbox_value = [float(x) for x in bbox]
                    elif len(bbox) >= 4 and all(isinstance(x, (list, tuple)) for x in bbox):
                        xs = [float(point[0]) for point in bbox]
                        ys = [float(point[1]) for point in bbox]
                        bbox_value = [min(xs), min(ys), max(xs), max(ys)]
                except (TypeError, ValueError, IndexError):
                    bbox_value = None
            text = normalize_line_text(direct)
            if text:
                collected.append((text, confidence, bbox_value))
        for key, nested in value.items():
            if key in {"text", "transcription", "confidence", "score", "text_region", "bbox", "box", "html"}:
                continue
            collected.extend(_text_from_nested(nested))
        return collected
    if isinstance(value, (list, tuple)):
        # PaddleOCR classic tuple: [quad, (text, score)]
        if (
            len(value) == 2
            and isinstance(value[1], (list, tuple))
            and len(value[1]) >= 1
            and isinstance(value[1][0], str)
        ):
            text = normalize_line_text(value[1][0])
            confidence = None
            if len(value[1]) > 1:
                try:
                    confidence = float(value[1][1])
                except (TypeError, ValueError):
                    confidence = None
            if text:
                collected.append((text, confidence, None))
            return collected
        for item in value:
            collected.extend(_text_from_nested(item))
    return collected


def parse_ppstructure_result(result: Any, page: int) -> List[TextLine]:
    lines: List[TextLine] = []
    regions: Sequence[Any] = result if isinstance(result, (list, tuple)) else [result]
    for region in regions:
        if not isinstance(region, dict):
            for text, confidence, bbox in _text_from_nested(region):
                lines.append(TextLine(text=text, page=page, bbox=bbox, confidence=confidence, block_type="ocr"))
            continue
        block_type = str(region.get("type") or "ocr")
        region_bbox = _bbox_from_region(region)
        payload = region.get("res", region)
        nested = _text_from_nested(payload)
        for text, confidence, text_bbox in nested:
            lines.append(
                TextLine(
                    text=text,
                    page=page,
                    bbox=text_bbox or region_bbox,
                    confidence=confidence,
                    block_type=block_type,
                )
            )
    # PPStructure may return repeated text through nested dicts. Preserve order.
    deduplicated: List[TextLine] = []
    seen: set[Tuple[int, str, str]] = set()
    for line in lines:
        marker = (page, line.block_type or "", line.text)
        if marker not in seen:
            seen.add(marker)
            deduplicated.append(line)
    return deduplicated


def _pixmap_to_bgr(page: fitz.Page, dpi: int) -> Any:
    import cv2
    import numpy as np

    pix = page.get_pixmap(dpi=dpi, alpha=False)
    array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        return cv2.cvtColor(array, cv2.COLOR_RGBA2BGR)
    return cv2.cvtColor(array, cv2.COLOR_RGB2BGR)


def _pdf_text_lines(page: fitz.Page, page_number: int) -> List[TextLine]:
    lines: List[TextLine] = []
    try:
        page_dict = page.get_text("dict")
        for block in page_dict.get("blocks", []):
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = normalize_line_text("".join(str(span.get("text") or "") for span in spans))
                if not text:
                    continue
                bbox = line.get("bbox") or block.get("bbox")
                bbox_value = [float(item) for item in bbox] if isinstance(bbox, (list, tuple)) and len(bbox) == 4 else None
                lines.append(TextLine(text=text, page=page_number, bbox=bbox_value, block_type="pdf_text"))
    except Exception:
        LOGGER.warning("Structured PDF text extraction failed; falling back to plain text", exc_info=True)
    if not lines:
        for text_line in (page.get_text("text") or "").splitlines():
            text = normalize_line_text(text_line)
            if text:
                lines.append(TextLine(text=text, page=page_number, block_type="pdf_text"))
    return lines


def source_from_pdf(path: Path, original_name: str, remove_seal: bool) -> SourceDocument:
    document = fitz.open(path)
    all_lines: List[TextLine] = []
    methods: List[str] = []
    warnings: List[str] = []
    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            raw_text = page.get_text("text") or ""
            normalized_chars = len("".join(raw_text.split()))
            if normalized_chars >= PDF_TEXT_PAGE_MIN_CHARS:
                methods.append("pdf_text")
                all_lines.extend(_pdf_text_lines(page, page_index + 1))
            else:
                methods.append("ocr")
                image = _pixmap_to_bgr(page, PDF_RENDER_DPI)
                if remove_seal:
                    image = redact_red_seal(image)
                recognized = OCR_ENGINE.recognize(image, page=page_index + 1)
                if recognized:
                    all_lines.extend(recognized)
                elif raw_text.strip():
                    # Never throw away a weak but nonempty text layer.
                    warnings.append(f"第 {page_index + 1} 页 OCR 无结果，回退到稀疏 PDF 文本层。")
                    for line in raw_text.splitlines():
                        text = normalize_line_text(line)
                        if text:
                            all_lines.append(TextLine(text=text, page=page_index + 1, block_type="pdf_text_fallback"))
        method_set = set(methods)
        if method_set == {"pdf_text"}:
            method = "pdf_text"
        elif method_set == {"ocr"}:
            method = "pdf_ocr"
        else:
            method = "pdf_hybrid"
        return SourceDocument(
            filename=original_name,
            extension=path.suffix.lower(),
            text="\n".join(line.text for line in all_lines),
            lines=all_lines,
            extraction_method=method,
            page_count=document.page_count,
            sha256=sha256_file(path),
            warnings=warnings,
        )
    finally:
        document.close()


def source_from_image(path: Path, original_name: str, remove_seal: bool) -> SourceDocument:
    import cv2
    import numpy as np

    data = np.fromfile(str(path), dtype=np.uint8)
    image = cv2.imdecode(data, cv2.IMREAD_COLOR)
    if image is None:
        raise ValueError("无法解码图片")
    if remove_seal:
        image = redact_red_seal(image)
    lines = OCR_ENGINE.recognize(image, page=1)
    return SourceDocument(
        filename=original_name,
        extension=path.suffix.lower(),
        text="\n".join(line.text for line in lines),
        lines=lines,
        extraction_method="image_ocr",
        page_count=1,
        sha256=sha256_file(path),
    )


def source_from_docx(path: Path, original_name: str) -> SourceDocument:
    from docx import Document

    document = Document(path)
    lines: List[TextLine] = []
    paragraph_index = 0
    for paragraph in document.paragraphs:
        paragraph_index += 1
        text = normalize_line_text(paragraph.text)
        if text:
            lines.append(TextLine(text=text, paragraph=paragraph_index, block_type="paragraph"))
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cell_texts = [normalize_line_text(cell.text) for cell in row.cells]
            cell_texts = [text for text in cell_texts if text]
            if cell_texts:
                lines.append(
                    TextLine(
                        text=" | ".join(cell_texts),
                        paragraph=paragraph_index + table_index,
                        block_type=f"table_row_{row_index}",
                    )
                )
    return SourceDocument(
        filename=original_name,
        extension=path.suffix.lower(),
        text="\n".join(line.text for line in lines),
        lines=lines,
        extraction_method="docx_text",
        page_count=None,
        sha256=sha256_file(path),
    )


def load_source(path: Path, original_name: str, remove_seal: bool = False) -> SourceDocument:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return source_from_pdf(path, original_name, remove_seal)
    if extension == ".docx":
        return source_from_docx(path, original_name)
    if extension in ALLOWED_EXTENSIONS:
        return source_from_image(path, original_name, remove_seal)
    raise ValueError(f"不支持的文件格式：{extension}")


# ---------------------------------------------------------------------------
# HTTP helpers and routes
# ---------------------------------------------------------------------------


def request_bool(name: str, default: bool = False) -> bool:
    value = request.form.get(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on", "y"}


def check_api_key() -> Optional[Tuple[Any, int]]:
    if not API_KEY:
        return None
    supplied = request.headers.get("X-API-Key", "")
    if supplied != API_KEY:
        return jsonify({"success": False, "error": {"code": "UNAUTHORIZED", "message": "无效的 API Key"}}), 401
    return None


def error_response(code: str, message: str, status: int, request_id: Optional[str] = None) -> Tuple[Any, int]:
    payload: Dict[str, Any] = {"success": False, "error": {"code": code, "message": message}}
    if request_id:
        payload["request_id"] = request_id
    return jsonify(payload), status


@app.errorhandler(RequestEntityTooLarge)
def handle_too_large(_: RequestEntityTooLarge) -> Tuple[Any, int]:
    return error_response("PAYLOAD_TOO_LARGE", f"请求体超过 {MAX_CONTENT_MB} MB 限制", 413)


@app.get("/api/v1/health")
def health() -> Tuple[Any, int]:
    auth_error = check_api_key()
    if auth_error:
        return auth_error
    return jsonify(
        {
            "success": True,
            "service": "legal-document-ocr",
            "version": "1.0.0",
            "ocr_engine": {
                "loaded": OCR_ENGINE.loaded,
                "load_error": OCR_ENGINE.load_error,
                "use_gpu": env_bool("OCR_USE_GPU", True),
            },
            "limits": {"max_content_mb": MAX_CONTENT_MB, "max_files": MAX_FILES},
        }
    ), 200


@app.get("/api/v1/schemas")
def schemas() -> Tuple[Any, int]:
    auth_error = check_api_key()
    if auth_error:
        return auth_error
    return jsonify({"success": True, "schema": api_schema()}), 200


def _extract_one(path: Path, original_name: str, hint: Optional[str], remove_seal: bool, include_raw_text: bool) -> Dict[str, Any]:
    source = load_source(path, original_name, remove_seal=remove_seal)
    result = extract_legal_document(source, document_type_hint=hint)
    if include_raw_text:
        result["raw_text"] = source.text
    return result


def _extract_uploaded_files() -> Tuple[Any, int]:
    auth_error = check_api_key()
    if auth_error:
        return auth_error

    request_id = str(uuid.uuid4())
    uploaded = request.files.getlist("files") or request.files.getlist("file")
    if not uploaded:
        return error_response("NO_FILES", "请使用 multipart/form-data 上传 files 字段", 400, request_id)
    if len(uploaded) > MAX_FILES:
        return error_response("TOO_MANY_FILES", f"单次最多上传 {MAX_FILES} 个文件", 400, request_id)

    hint = (request.form.get("document_type") or "").strip() or None
    remove_seal = request_bool("remove_red_seal", False)
    include_raw_text = request_bool("include_raw_text", False)
    documents: List[Dict[str, Any]] = []
    errors: List[Dict[str, Any]] = []

    with tempfile.TemporaryDirectory(prefix="legal_ocr_") as temp_dir:
        temp_root = Path(temp_dir)
        for index, storage in enumerate(uploaded):
            original_name = storage.filename or f"upload_{index + 1}"
            extension = Path(original_name).suffix.lower()
            if extension not in ALLOWED_EXTENSIONS:
                errors.append(
                    {
                        "filename": original_name,
                        "code": "UNSUPPORTED_FILE_TYPE",
                        "message": f"不支持的文件格式：{extension or '无扩展名'}",
                    }
                )
                continue
            safe_stem = secure_filename(Path(original_name).stem) or f"upload_{index + 1}"
            temp_path = temp_root / f"{uuid.uuid4().hex}_{safe_stem}{extension}"
            storage.save(temp_path)
            try:
                documents.append(_extract_one(temp_path, original_name, hint, remove_seal, include_raw_text))
            except Exception as exc:
                LOGGER.exception("Document extraction failed: %s", original_name)
                errors.append(
                    {
                        "filename": original_name,
                        "code": "EXTRACTION_FAILED",
                        "message": str(exc),
                    }
                )

    status = 200 if documents else 422
    return jsonify(
        {
            "success": bool(documents),
            "request_id": request_id,
            "documents": documents,
            "errors": errors,
            "summary": {
                "received": len(uploaded),
                "succeeded": len(documents),
                "failed": len(errors),
            },
        }
    ), status


@app.post("/api/v1/ocr/extract")
def extract_uploads() -> Tuple[Any, int]:
    return _extract_uploaded_files()


# Backward-compatible route for the former experimental frontend.
@app.post("/autofill_from_uploads")
def legacy_extract_uploads() -> Tuple[Any, int]:
    return _extract_uploaded_files()


@app.post("/api/v1/ocr/extract-text")
def extract_text() -> Tuple[Any, int]:
    auth_error = check_api_key()
    if auth_error:
        return auth_error
    payload = request.get_json(silent=True) or {}
    text = str(payload.get("text") or "").strip()
    if not text:
        return error_response("EMPTY_TEXT", "text 不能为空", 400)
    filename = str(payload.get("filename") or "inline.txt")
    hint = payload.get("document_type")
    lines = [TextLine(text=line, paragraph=index + 1, block_type="inline_text") for index, line in enumerate(text.splitlines()) if line.strip()]
    source = SourceDocument(
        filename=filename,
        extension=".txt",
        text=text,
        lines=lines,
        extraction_method="inline_text",
        sha256=hashlib.sha256(text.encode("utf-8")).hexdigest(),
    )
    return jsonify({"success": True, "document": extract_legal_document(source, document_type_hint=hint)}), 200


@app.get("/")
def index() -> Tuple[Any, int]:
    return jsonify(
        {
            "service": "legal-document-ocr",
            "version": "1.0.0",
            "endpoints": {
                "health": "GET /api/v1/health",
                "schemas": "GET /api/v1/schemas",
                "extract": "POST /api/v1/ocr/extract",
            },
        }
    ), 200


if __name__ == "__main__":
    app.run(host=APP_HOST, port=APP_PORT, threaded=True)
