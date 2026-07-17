# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import time
import re
from copy import deepcopy

from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger
logger = get_logger()


def add_textbox_to_paragraph(paragraph, text, left_emu, top_emu, width_emu, height_emu, 
                             bold=False, italic=False, size=10, color_rgb=None, center=False, box_id=1):
    """
    在段落中添加绝对定位的文本框
    left_emu, top_emu, width_emu, height_emu: EMU单位（1英寸=914400 EMU）
    """
    # 处理文本（转义特殊字符）
    text = str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
    
    # 构建文本框XML
    textbox_xml = f'''
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
         xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
         xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
         xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <w:drawing>
            <wp:anchor distT="0" distB="0" distL="114300" distR="114300" 
                       simplePos="0" relativeHeight="251658240" behindDoc="0" 
                       locked="0" layoutInCell="1" allowOverlap="1">
                <wp:simplePos x="0" y="0"/>
                <wp:positionH relativeFrom="page">
                    <wp:posOffset>{left_emu}</wp:posOffset>
                </wp:positionH>
                <wp:positionV relativeFrom="page">
                    <wp:posOffset>{top_emu}</wp:posOffset>
                </wp:positionV>
                <wp:extent cx="{width_emu}" cy="{height_emu}"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:wrapNone/>
                <wp:docPr id="{box_id}" name="TextBox {box_id}"/>
                <wp:cNvGraphicFramePr/>
                <a:graphic>
                    <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                        <wps:txbx>
                            <w:txbxContent>
                                <w:p>
                                    <w:pPr>
                                        <w:jc w:val="{'center' if center else 'left'}"/>
                                    </w:pPr>
                                    <w:r>
                                        <w:rPr>
                                            <w:b w:val="{'1' if bold else '0'}"/>
                                            <w:i w:val="{'1' if italic else '0'}"/>
                                            <w:sz w:val="{int(size * 2)}"/>
                                        </w:rPr>
                                        <w:t>{text}</w:t>
                                    </w:r>
                                </w:p>
                            </w:txbxContent>
                        </wps:txbx>
                        <wps:bodyPr rot="0" vert="horz" anchor="ctr" anchorCtr="0"/>
                    </a:graphicData>
                </a:graphic>
            </wp:anchor>
        </w:drawing> 
    </w:r>
    '''
    
    try:
        # 将XML添加到段落
        run_element = parse_xml(textbox_xml)
        paragraph._element.append(run_element)
    except Exception as e:
        logger.warning(f'Failed to add textbox: {e}')


def convert_info_docx_with_textboxes(all_res_list, save_folder, img_name):
    """
    直接用Word文本框实现绝对定位，保留原PDF布局
    all_res_list: 包含 {'img': img, 'res': res, 'page_index': page_index} 的列表
    """
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)
    
    # 设置页面大小为A4
    section = doc.sections[0]
    section.page_height = shared.Inches(11)
    section.page_width = shared.Inches(8.5)
    section.top_margin = shared.Inches(0.5)
    section.bottom_margin = shared.Inches(0.5)
    section.left_margin = shared.Inches(0.5)
    section.right_margin = shared.Inches(0.5)
    
    textbox_id = 1
    
    for page_idx, page_data in enumerate(all_res_list):
        res = page_data['res']
        page_index = page_data['page_index']
        
        # 添加分页符（除了第一页）
        if page_idx > 0:
            doc.add_page_break()
        
        # 按y坐标排序
        sorted_res = sorted(res, key=lambda r: (r['bbox'][1], r['bbox'][0]))
        
        # 创建一个空段落作为容器
        page_paragraph = doc.add_paragraph()
        
        # 遍历每个region
        for region in sorted_res:
            if len(region['res']) == 0:
                continue
            
            bbox = region['bbox']  # [x1, y1, x2, y2]
            region_type = region['type'].lower()
            
            # 转换为EMU单位（1英寸=914400 EMU）
            x_emu = int(bbox[0] * 914400 / 72)
            y_emu = int(bbox[1] * 914400 / 72)
            width_emu = int((bbox[2] - bbox[0]) * 914400 / 72)
            height_emu = int((bbox[3] - bbox[1]) * 914400 / 72)
            
            # 确保最小尺寸
            width_emu = max(width_emu, 914400)  # 至少1英寸
            height_emu = max(height_emu, 200000)  # 至少0.2英寸
            
            if region_type == 'header':
                # header：加粗、3号字体（16pt）、居中
                for line in region['res']:
                    text = line['text'] if isinstance(line, dict) else line
                    add_textbox_to_paragraph(page_paragraph, text, x_emu, y_emu, width_emu, height_emu,
                                           bold=True, size=16, center=True, box_id=textbox_id)
                    textbox_id += 1
            
            elif region_type == 'title':
                # title：加粗、大字体
                for line in region['res']:
                    text = line['text'] if isinstance(line, dict) else line
                    add_textbox_to_paragraph(page_paragraph, text, x_emu, y_emu, width_emu, height_emu,
                                           bold=True, size=14, box_id=textbox_id)
                    textbox_id += 1
            
            elif region_type == 'footer':
                # footer：居中、小字体
                for line in region['res']:
                    text = line['text'] if isinstance(line, dict) else line
                    add_textbox_to_paragraph(page_paragraph, text, x_emu, y_emu, width_emu, height_emu,
                                           size=10, center=True, box_id=textbox_id)
                    textbox_id += 1
            
            elif region_type == 'figure':
                # 图片：如果有图片路径则添加占位符
                add_textbox_to_paragraph(page_paragraph, '[图片]', x_emu, y_emu, width_emu, height_emu,
                                       size=10, center=True, box_id=textbox_id)
                textbox_id += 1
            
            elif region_type == 'table':
                # 表格：显示为文本
                if isinstance(region['res'], dict) and 'html' in region['res']:
                    html_table = region['res']['html']
                    # 简单地去除HTML标签显示文本
                    text_content = re.sub(r'<[^>]+>', '', html_table)
                    add_textbox_to_paragraph(page_paragraph, text_content, x_emu, y_emu, width_emu, height_emu,
                                           size=10, box_id=textbox_id)
                    textbox_id += 1
            
            else:
                # 默认文本
                for line in region['res']:
                    text = line['text'] if isinstance(line, dict) else line
                    add_textbox_to_paragraph(page_paragraph, text, x_emu, y_emu, width_emu, height_emu,
                                           size=10, box_id=textbox_id)
                    textbox_id += 1
    
    # 保存文档（带重试机制）
    max_retries = 3
    for attempt in range(max_retries):
        try:
            # 检查文件是否被锁定
            if os.path.exists(docx_path):
                try:
                    with open(docx_path, 'ab') as f:
                        pass
                except (IOError, OSError) as e:
                    # 文件被锁定，生成备份
                    logger.warning(f'File {docx_path} is locked, saving as backup')
                    timestamp = int(time.time() * 1000) % 100000
                    docx_path = os.path.join(save_folder, f'{img_name}_{timestamp}_ocr.docx')
                    logger.info(f'Using backup path: {docx_path}')
            
            doc.save(docx_path)
            logger.info(f'Document saved successfully to {docx_path}')
            break
        
        except (IOError, OSError, PermissionError) as e:
            if attempt < max_retries - 1:
                logger.warning(f'Failed to save (attempt {attempt + 1}/{max_retries}): {e}, retrying...')
                time.sleep(1)
            else:
                # 最后尝试保存到备份位置
                timestamp = int(time.time() * 1000) % 100000
                backup_path = os.path.join(save_folder, f'{img_name}_{timestamp}_ocr.docx')
                try:
                    doc.save(backup_path)
                    logger.info(f'Failed to save to original path, saved to backup: {backup_path}')
                except Exception as final_error:
                    logger.error(f'Failed to save document: {final_error}')
                    raise


def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)

    flag = 1
    for i, region in enumerate(res):
        if len(region['res']) == 0:
            continue
        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        if region['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(region['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region['type'].lower() == 'title':
            doc.add_heading(region['res'][0]['text'])
        elif region['type'].lower() == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            try:
                parser.handle_table(region['res']['html'], doc)
            except Exception as e:
                logger.warning(f"Failed to parse table, using text fallback: {e}")
                paragraph = doc.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                if isinstance(region['res'], dict) and 'html' in region['res']:
                    html_content = region['res']['html']
                    text_content = re.sub(r'<[^>]+>', '', html_content)
                    paragraph.add_run(text_content)
                else:
                    paragraph.add_run("[表格识别失败，请查看生成的图片]")
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)
    logger.info('docx save to {}'.format(docx_path))


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res


def convert_info_docx_multi_page(all_res_list, save_folder, img_name):
    """
    将多页的识别结果合并到一个Word文档中
    直接用Word文本框实现绝对定位，完全保留原PDF布局
    """
    logger.info('Starting textbox-based document generation...')
    convert_info_docx_with_textboxes(all_res_list, save_folder, img_name)
    logger.info('Document generation completed')
