# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import time
import subprocess
import re
from copy import deepcopy
from difflib import SequenceMatcher
import cv2

from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger
logger = get_logger()


def group_text_by_lines(text_results, y_threshold=20):
    """
    极简版：对比每一个item的text_region的中心y坐标，
    如果相差 < y_threshold 则合并成一行
    """
    if not text_results:
        return []
        
    items = []
    for idx, item in enumerate(text_results):
        # 容错兜底：绝对不能丢弃文本数据
        if not isinstance(item, dict):
            items.append({'text': str(item), 'x': 0, 'x2': 100, 'y': idx * 1000})
            continue
            
        text = item.get('text', '')
        tr = item.get('text_region', None)
        
        if tr is None:
            items.append({'text': text, 'x': 0, 'x2': 100, 'y': idx * 1000})
            continue
            
        try:
            # 兼容 numpy array 输出的情况
            if hasattr(tr, 'tolist'):
                tr = tr.tolist()
                
            # 获取坐标，用 float() 强转防止 np.float32 坑人
            if len(tr) == 4 and not isinstance(tr[0], (list, tuple)):
                x1, y1, x2, y2 = float(tr[0]), float(tr[1]), float(tr[2]), float(tr[3])
            elif len(tr) >= 4 and isinstance(tr[0], (list, tuple)):
                x1 = min(float(p[0]) for p in tr)
                x2 = max(float(p[0]) for p in tr)
                y1 = min(float(p[1]) for p in tr)
                y2 = max(float(p[1]) for p in tr)
            else:
                x1, y1, x2, y2 = 0, idx * 1000, 100, idx * 1000 + 20
        except Exception:
            # 任何报错都不应当丢失文字内容
            x1, y1, x2, y2 = 0, idx * 1000, 100, idx * 1000 + 20
            
        cy = (y1 + y2) / 2.0
        items.append({
            'text': text,
            'x': x1,
            'x2': x2,
            'y': cy
        })

    if not items:
        return []

    # 先整体按y坐标排序
    items.sort(key=lambda i: i['y'])
    
    lines = []
    current_line = [items[0]]
    
    for item in items[1:]:
        # 互相对比，只要和当前行最后加入的元素y差距在 y_threshold 范围内，就是同行（直接用用户的硬性规则）
        if abs(item['y'] - current_line[-1]['y']) < y_threshold:
            current_line.append(item)
        else:
            current_line.sort(key=lambda i: i['x'])
            lines.append(current_line)
            current_line = [item]
            
    if current_line:
        current_line.sort(key=lambda i: i['x'])
        lines.append(current_line)
        
    return lines


def merge_line_texts(line_items, use_spacing=True, char_width=10):
    """
    合并同一行的文本，保持适当的间距
    
    Args:
        line_items: 同一行的文本项列表
        use_spacing: 是否根据x坐标添加空格
        char_width: 估算的每个字符宽度（像素）
    
    Returns:
        合并后的行文本
    """
    if not line_items:
        return ""
    
    if len(line_items) == 1:
        return line_items[0]['text']
    
    result_parts = []
    for i, item in enumerate(line_items):
        if i == 0:
            result_parts.append(item['text'])
        else:
            # 计算与前一个文本框的间距
            gap = item['x'] - line_items[i-1]['x2']
            
            if use_spacing and gap > char_width * 2:
                # 较大间距，添加多个空格或制表符
                num_spaces = max(1, int(gap / char_width / 2))
                result_parts.append(' ' * num_spaces)
            else:
                # 普通间距，添加一个空格
                result_parts.append(' ')
            
            result_parts.append(item['text'])
    
    return ''.join(result_parts)


def _escape_html(text):
    """转义HTML特殊字符"""
    if not isinstance(text, str):
        text = str(text)
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&#39;')
    return text


def generate_html_from_regions(all_res_list, save_folder, img_name):
    """
    从layout结果生成HTML，使用绝对定位保留原始布局
    all_res_list: 包含 {'img': img, 'res': res, 'page_index': page_index} 的列表
    """
    html_content = '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {
            margin: 0;
            padding: 20px;
            font-family: 'Times New Roman', 'SimSun', serif;
            background: white;
        }
        .page {
            position: relative;
            width: 8.5in;
            height: 11in;
            margin: 20px auto;
            padding: 0.5in;
            background: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            page-break-after: always;
        }
        .region {
            position: absolute;
            box-sizing: border-box;
        }
        .region-header {
            text-align: center;
            font-weight: bold;
            font-size: 18pt;
        }
        .region-title {
            font-weight: bold;
            font-size: 16pt;
            margin-bottom: 10px;
        }
        .region-text {
            font-size: 11pt;
            line-height: 1.5;
            text-indent: 0.5in;
        }
        .region-footer {
            font-size: 10pt;
            text-align: center;
        }
        .region-table {
            font-size: 10pt;
        }
        .region-figure {
            text-align: center;
        }
        .region-figure img {
            max-width: 90%;
            max-height: 90%;
        }
        table {
            border-collapse: collapse;
            width: 100%;
        }
        table, th, td {
            border: 1px solid #000;
        }
        th, td {
            padding: 5px;
            text-align: left;
        }
    </style>
</head>
<body>
'''
    
    # 遍历每一页
    for page_idx, page_data in enumerate(all_res_list):
        res = page_data['res']
        page_index = page_data['page_index']
        
        html_content += f'    <div class="page" id="page-{page_index}">\n'
        html_content += f'        <div style="font-size: 8pt; color: #999; margin-bottom: 10px;">页 {page_index + 1}</div>\n'
        
        # 按y坐标排序
        sorted_res = sorted(res, key=lambda r: (r['bbox'][1], r['bbox'][0]))
        
        # 遍历每个region
        for region in sorted_res:
            if len(region['res']) == 0:
                continue
            
            bbox = region['bbox']
            # bbox格式: [x1, y1, x2, y2]
            # 转换为HTML坐标（像素到英寸）
            x = bbox[0] / 72  # 72 DPI
            y = bbox[1] / 72
            width = (bbox[2] - bbox[0]) / 72
            height = (bbox[3] - bbox[1]) / 72
            
            region_type = region['type'].lower()
            
            # 开始region div
            html_content += f'        <div class="region region-{region_type}" style="left: {x}in; top: {y}in; width: {width}in; height: {height}in;">\n'
            
            if region_type == 'header':
                for line in region['res']:
                    html_content += f'            <div class="region-header">{_escape_html(line["text"])}</div>\n'
            
            elif region_type == 'title':
                for line in region['res']:
                    html_content += f'            <div class="region-title">{_escape_html(line["text"])}</div>\n'
            
            elif region_type == 'footer':
                for line in region['res']:
                    html_content += f'            <div class="region-footer">{_escape_html(line["text"])}</div>\n'
            
            elif region_type == 'figure':
                # 图片处理
                img_idx = region['img_idx']
                img_path = region.get('img_path', '')
                if img_path:
                    html_content += f'            <div class="region-figure"><img src="{img_path}" alt="figure"></div>\n'
            
            elif region_type == 'table':
                if isinstance(region['res'], dict) and 'html' in region['res']:
                    html_table = region['res']['html']
                    html_content += f'            <div class="region-table">{html_table}</div>\n'
            
            else:  # 默认文本
                for line_idx, line in enumerate(region['res']):
                    text = line['text'] if isinstance(line, dict) else line
                    html_content += f'            <div class="region-text">{_escape_html(text)}</div>\n'
            
            html_content += '        </div>\n'
        
        html_content += '    </div>\n'
    
    html_content += '''</body>
</html>
'''
    return html_content


def convert_html_to_docx(html_path, docx_path):
    """
    使用pandoc将HTML转换为DOCX
    """
    try:
        # 检查pandoc是否安装
        subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
        
        # 使用pandoc转换
        cmd = [
            'pandoc',
            html_path,
            '-o', docx_path,
            '--from=html',
            '--to=docx',
            '-V', 'margin-left=0.5in',
            '-V', 'margin-right=0.5in',
            '-V', 'margin-top=0.5in',
            '-V', 'margin-bottom=0.5in'
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            logger.info(f'Successfully converted HTML to DOCX using pandoc: {docx_path}')
            return True
        else:
            logger.warning(f'Pandoc conversion failed: {result.stderr}')
            return False
    
    except FileNotFoundError:
        logger.warning('Pandoc not found, will use fallback method')
        return False
    except Exception as e:
        logger.warning(f'Error during pandoc conversion: {e}')
        return False


def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)

    # 预扫描：判断是否为合同电子签章申请函（相似度>=80%）
    _CONTRACT_ESIGN_TITLE = '合同电子签章申请函'
    is_contract_esign = False
    for _r in res:
        if (_r['type'].lower() == 'title' or _r['type'].lower() == 'header') and _r['res']:
            _title_text = ' '.join(
                item.get('text', '') if isinstance(item, dict) else str(item)
                for item in _r['res']
            ).strip()
            if SequenceMatcher(None, _title_text, _CONTRACT_ESIGN_TITLE).ratio() >= 0.8:
                is_contract_esign = True
                logger.info(f'[contract_esign] 检测到合同电子签章申请函，将启用表格对齐与修正')
                break

    flag = 1
    for i, region in enumerate(res):
        if len(region['res']) == 0:
            continue
        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        if region['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(region['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region['type'].lower() == 'title':
            doc.add_heading(region['res'][0]['text'])
        elif region['type'].lower() == 'table':
            print(1111111111111)
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            try:
                parser.handle_table(region['res']['html'], doc)
                print('debugggg',region['res']['html'])
                if is_contract_esign and doc.tables:
                    _repair_contract_esign_docx_table(doc.tables[-1])
                    _fix_name_phone_in_docx_table(doc.tables[-1])
            except Exception as e: 
                import traceback
                logger.warning(f"Failed to parse table, using text fallback: {e}\n{traceback.format_exc()}")
                # 表格处理失败时，改为插入文本内容
                paragraph = doc.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                if isinstance(region['res'], dict) and 'html' in region['res']:
                    # 尝试从 HTML 中提取文本
                    html_content = region['res']['html']
                    # 简单的 HTML 标签移除
                    import re
                    text_content = re.sub(r'<[^>]+>', '', html_content)
                    paragraph.add_run(text_content)
                else:
                    # 如果是其他格式，插入提示信息
                    paragraph.add_run("[表格识别失败，请查看生成的图片]")
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)
    logger.info('docx save to {}'.format(docx_path))


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res


    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res


# 合同电子签章表格 Schema 修复及对齐配置与实现
# 测试 1：后缀污染
# 输入：
#     key = `收件人地址：号`
#     value = `珠海市斗门区斗门镇八甲村旧赤水坑队1`
# expected_key = `收件人地址`
# 输出：
#     key = `收件人地址：`
#     value = `珠海市斗门区斗门镇八甲村旧赤水坑队1号`
#
# 测试 2：前缀污染
# 输入：
#     prev_key = `银行开户行：`
#     prev_value = `中国农业银行股份有限公司舟山南珍支`
#     key = `行 银行账号：`
#     value = `19425101040090774`
# expected_key = `银行账号`
# 输出：
#     prev_value = `中国农业银行股份有限公司舟山南珍支行`
#     key = `银行账号：`
#     value = `19425101040090774`
#
# 测试 3：无“收件人地址”的 schema
# 输入左列顺序：
#     `收件人姓名：`
#     `收件人电话：`
#     `收件人电子邮箱：`
#     `名称：`
#     `纳税人识别号：`
#     `地址：`
#     `电话：`
#     `银行开户行：`
#     `银行账号：`
# 输出应保持这个 schema，不要强行插入或补出 `收件人地址：`。

CONTRACT_ESIGN_SCHEMA_NO_ADDRESS = [
    "收件人姓名",
    "收件人电话",
    "收件人电子邮箱",
    "名称",
    "纳税人识别号",
    "地址",
    "电话",
    "银行开户行",
    "银行账号",
]

CONTRACT_ESIGN_SCHEMA_WITH_ADDRESS = [
    "收件人姓名",
    "收件人电话",
    "收件人地址",
    "收件人电子邮箱",
    "名称",
    "纳税人识别号",
    "地址",
    "电话",
    "银行开户行",
    "银行账号",
]

CONTRACT_ESIGN_SCHEMAS = [
    CONTRACT_ESIGN_SCHEMA_WITH_ADDRESS,
    CONTRACT_ESIGN_SCHEMA_NO_ADDRESS,
]


def _get_unique_row_cells(table):
    """
    输入 python-docx 的 table。
    返回二维列表，每一行的唯一 cell（逐行内部去重，避免全表级去重在合并单元格场景下错误跳过后续行）。
    示例结构：[[cell00, cell01], [cell10, cell11], ...]
    """
    all_rows_cells = []
    for row in table.rows:
        row_cells = []
        seen_tc = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id not in seen_tc:
                seen_tc.add(tc_id)
                row_cells.append(cell)
        all_rows_cells.append(row_cells)
    return all_rows_cells


def _get_cell_text(cell):
    """
    返回 cell.text 的清洗结果。
    去掉首尾空白，把连续空白、换行、制表符压缩为一个空格。
    """
    if cell is None:
        return ""
    text = cell.text
    if not isinstance(text, str):
        text = str(text)
    text = re.sub(r'[\s\n\t]+', ' ', text)
    return text.strip()


def _replace_cell_text(cell, text):
    """
    清空 cell 内已有 run 的 text，然后写入新 text。
    不破坏段落结构，没有 paragraph 则新建。
    """
    if not cell:
        return
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if not cell.paragraphs:
        cell.add_paragraph()
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _normalize_key_for_match(text):
    """
    用于标题匹配。
    去掉空格、换行、制表符。
    统一中文冒号和英文冒号。
    去掉末尾冒号及常见干扰符号。
    """
    if not text:
        return ""
    text = re.sub(r'\s+', '', text)
    text = text.replace(':', '：')
    while text.endswith('：'):
        text = text[:-1]
    # 去除常见干扰符
    text = re.sub(r'[\*\-_．\.#@、，,（）\(\)]', '', text)
    return text


def _ensure_cn_colon(text):
    """
    保证标题以中文冒号 `：` 结尾。
    如果已有尾部英文或中文冒号，统一为中文冒号，空字符串不加。
    """
    if not text:
        return ""
    text = text.strip()
    while text.endswith(':') or text.endswith('：'):
        text = text[:-1].strip()
    if text:
        return text + "："
    return ""


def _schema_key_similarity(raw_text, schema_key):
    """
    使用 _normalize_key_for_match 后的文本做相似度。
    若 key_norm 是 raw_norm 子串加 0.25；若反之加 0.1；最后最高 1.0。
    """
    raw_norm = _normalize_key_for_match(raw_text)
    key_norm = _normalize_key_for_match(schema_key)
    if not raw_norm or not key_norm:
        return 0.0
    
    score = SequenceMatcher(None, raw_norm, key_norm).ratio()
    if len(key_norm) > 0 and key_norm in raw_norm:
        score += 0.25
    elif len(raw_norm) > 0 and raw_norm in key_norm:
        score += 0.1
    return min(1.0, score)


def _best_schema_key_match(raw_text, schema):
    """
    输入一个 raw left cell 文本和一个 schema。
    返回 (best_key, best_score)。
    """
    best_key = None
    best_score = -1.0
    for key in schema:
        score = _schema_key_similarity(raw_text, key)
        if score > best_score:
            best_score = score
            best_key = key
    return best_key, best_score


def _split_suffix_after_colon(raw_left, expected_key, schema=None):
    """
    用于处理后缀污染。
    """
    if ':' in raw_left or '：' in raw_left:
        idx = raw_left.find(':')
        idx_cn = raw_left.find('：')
        if idx == -1:
            split_idx = idx_cn
        elif idx_cn == -1:
            split_idx = idx
        else:
            split_idx = min(idx, idx_cn)
            
        key_part = raw_left[:split_idx].strip()
        suffix = raw_left[split_idx+1:].strip()
        
        if 1 <= len(suffix) <= 4:
            sim_score = _schema_key_similarity(key_part, expected_key)
            similar_enough = (sim_score >= 0.55)
            
            if not similar_enough and schema:
                _, best_score = _best_schema_key_match(key_part, schema)
                if best_score >= 0.55:
                    similar_enough = True
            
            if similar_enough:
                return expected_key, suffix
    return None


def _previous_value_completion_score(prev_value, prefix):
    """
    判断 prefix 是否应该接回上一行右侧 value。
    """
    if not prefix:
        return 0.0
    combined = (prev_value or "") + prefix
    high_suffixes = ("支行", "分行", "银行", "公司", "有限公司", "号", "室", "楼", "层", "省", "市", "区", "县", "镇")
    medium_prefixes = ("行", "司", "号", "室", "楼", "层", "省", "市", "区")
    
    if combined.endswith(high_suffixes):
        return 1.0
    elif prefix in medium_prefixes and len(prefix) <= 2:
        return 0.6
    return 0.1


def _is_probable_value_fragment(fragment):
    """
    判断 suffix 或者是前缀是否可能是 value 的一部分。
    """
    if not fragment or len(fragment) > 4:
        return False
    if ':' in fragment or '：' in fragment:
        return False
        
    known_fragments = ("号", "室", "楼", "层", "行", "司", "公司", "支行", "分行", "省", "市", "区", "县", "镇")
    if fragment in known_fragments:
        return True
        
    if re.match(r'^[a-zA-Z0-9]+$', fragment):
        return True
        
    if re.match(r'^\d+[号室楼层区]$', fragment):
        return True
    return False


def _extract_kv_pairs_from_docx_table(table):
    """
    从 Word table 中抽取 key-value 行。
    每行至少两个 cell 时，取 row_cells[0] 作为 key_cell，row_cells[1] 作为 value_cell。
    """
    all_rows = _get_unique_row_cells(table)
    kv_pairs = []
    for i, row_cells in enumerate(all_rows):
        if len(row_cells) >= 2:
            key_cell = row_cells[0]
            value_cell = row_cells[1]
            key_text = _get_cell_text(key_cell)
            value_text = _get_cell_text(value_cell)
            kv_pairs.append({
                "row_index": i,
                "key_cell": key_cell,
                "value_cell": value_cell,
                "raw_key": key_text,
                "raw_value": value_text,
            })
    return kv_pairs


def _choose_contract_esign_schema(kv_pairs):
    """
    从两种合法 schema 中选择更适合当前表格的 schema。
    """
    if not kv_pairs:
        return CONTRACT_ESIGN_SCHEMA_WITH_ADDRESS
        
    score_with = 0.0
    score_no = 0.0
    
    has_address = False
    has_email = False
    has_phone = False
    
    for pair in kv_pairs:
        rk = pair["raw_key"]
        
        best_with_sim = max(_schema_key_similarity(rk, sk) for sk in CONTRACT_ESIGN_SCHEMA_WITH_ADDRESS)
        best_no_sim = max(_schema_key_similarity(rk, sk) for sk in CONTRACT_ESIGN_SCHEMA_NO_ADDRESS)
        
        score_with += best_with_sim
        score_no += best_no_sim
        
        if _schema_key_similarity(rk, "收件人地址") >= 0.7:
            has_address = True
        if _schema_key_similarity(rk, "收件人电子邮箱") >= 0.7:
            has_email = True
        if _schema_key_similarity(rk, "收件人电话") >= 0.7:
            has_phone = True
            
    if has_address:
        score_with += 5.0
    elif has_email and has_phone:
        score_no += 2.0
        
    if score_with >= score_no:
        return CONTRACT_ESIGN_SCHEMA_WITH_ADDRESS
    else:
        return CONTRACT_ESIGN_SCHEMA_NO_ADDRESS


def _align_kv_pairs_to_schema(kv_pairs, schema):
    """
    基于贪心+顺序约束对齐 kv_pair 字典键至 schema。
    """
    aligned_keys = []
    schema_pos = 0
    
    for pair in kv_pairs:
        rk = pair["raw_key"]
        
        best_suffix_key = None
        best_suffix_idx = -1
        best_suffix_score = -1.0
        
        for i in range(schema_pos, len(schema)):
            sk = schema[i]
            score = _schema_key_similarity(rk, sk)
            if score > best_suffix_score:
                best_suffix_score = score
                best_suffix_key = sk
                best_suffix_idx = i
                
        if best_suffix_score >= 0.55:
            aligned_keys.append(best_suffix_key)
            schema_pos = best_suffix_idx + 1
        else:
            # 全局兜底
            best_global_key = None
            best_global_idx = -1
            best_global_score = -1.0
            for i, sk in enumerate(schema):
                score = _schema_key_similarity(rk, sk)
                if score > best_global_score:
                    best_global_score = score
                    best_global_key = sk
                    best_global_idx = i
            if best_global_score >= 0.55:
                aligned_keys.append(best_global_key)
                schema_pos = best_global_idx + 1
            else:
                aligned_keys.append(None)
                
    return aligned_keys


def _repair_contract_esign_docx_table(table, max_contam_len=4):
    """
    主后处理函数，对表格左列标题进行强 schema 对齐与首尾污染修复。
    """
    try:
        kv_pairs = _extract_kv_pairs_from_docx_table(table)
        if not kv_pairs:
            return

        schema = _choose_contract_esign_schema(kv_pairs)
        aligned_keys = _align_kv_pairs_to_schema(kv_pairs, schema)

        for i, pair in enumerate(kv_pairs):
            expected_key = aligned_keys[i]
            if not expected_key:
                continue

            key_cell = pair["key_cell"]
            value_cell = pair["value_cell"]
            raw_left = _get_cell_text(key_cell)
            raw_right = value_cell.text.strip()
            prev_pair = kv_pairs[i - 1] if i > 0 else None

            # D1. 后缀污染优先
            split_res = _split_suffix_after_colon(raw_left, expected_key, schema)
            if split_res:
                fixed_key, suffix = split_res
                if _is_probable_value_fragment(suffix):
                    new_key_text = _ensure_cn_colon(expected_key)
                    new_value_text = raw_right + suffix
                    _replace_cell_text(key_cell, new_key_text)
                    _replace_cell_text(value_cell, new_value_text)
                    logger.info(f'[contract_esign_schema_repair] suffix moved: raw_key="{raw_left}", expected="{expected_key}", moved="{suffix}"')
                    continue

            # D2. 前缀污染
            # prefix_matched = False
            # if prev_pair:
            #     prev_value_cell = prev_pair["value_cell"]
            #     prev_value = prev_value_cell.text.strip()

            #     for prefix_len in range(1, max_contam_len + 1):
            #         if len(raw_left) > prefix_len:
            #             prefix = raw_left[:prefix_len].strip()
            #             rest_key = raw_left[prefix_len:].strip()

            #             if _schema_key_similarity(rest_key, expected_key) >= 0.65:
            #                 score = _previous_value_completion_score(prev_value, prefix)
            #                 if score >= 0.6:
            #                     new_prev_value = prev_value + prefix
            #                     new_key_text = _ensure_cn_colon(expected_key)
            #                     _replace_cell_text(prev_value_cell, new_prev_value)
            #                     _replace_cell_text(key_cell, new_key_text)
            #                     logger.info(f'[contract_esign_schema_repair] prefix moved: raw_key="{raw_left}", expected="{expected_key}", moved="{prefix}"')
            #                     prefix_matched = True
            #                     break
            #     if prefix_matched:
            #         continue
            # D2. 前缀污染：只在 raw_left 中存在空格分隔时才尝试修复
            # 例： "行 银行账号：" -> prefix="行", rest_key="银行账号："
            prefix_matched = False
            if prev_pair and re.search(r'\s+', raw_left):
                prev_value_cell = prev_pair["value_cell"]
                prev_value = prev_value_cell.text.strip()

                parts = re.split(r'\s+', raw_left.strip(), maxsplit=1)
                if len(parts) == 2:
                    prefix, rest_key = parts[0].strip(), parts[1].strip()

                    if 0 < len(prefix) <= max_contam_len:
                        if _schema_key_similarity(rest_key, expected_key) >= 0.65:
                            score = _previous_value_completion_score(prev_value, prefix)
                            if score >= 0.6:
                                new_prev_value = prev_value + prefix
                                new_key_text = _ensure_cn_colon(expected_key)
                                _replace_cell_text(prev_value_cell, new_prev_value)
                                _replace_cell_text(key_cell, new_key_text)
                                logger.info(f'[contract_esign_schema_repair] prefix moved: raw_key="{raw_left}", expected="{expected_key}", moved="{prefix}"')
                                prefix_matched = True

                if prefix_matched:
                    continue
            # D3. 普通标题归正
            raw_norm = _normalize_key_for_match(raw_left)
            expected_norm = _normalize_key_for_match(expected_key)

            sim = _schema_key_similarity(raw_left, expected_key)
            is_contained = (expected_norm in raw_norm) or (raw_norm in expected_norm and len(raw_norm) > 0)

            if sim >= 0.55 or is_contained:
                new_key_text = _ensure_cn_colon(expected_key)
                _replace_cell_text(key_cell, new_key_text)
                logger.info(f'[contract_esign_schema_repair] key canonicalized: raw_key="{raw_left}", expected="{expected_key}"')

    except Exception as e:
        logger.warning(f"[contract_esign_schema_repair] Error during contract e-sign table repair: {e}")


def _fix_name_phone_in_docx_table(table):
    """
    对 Word 表格后处理：若 "收件人姓名" 单元格尾部混入了电话号码，
    且对应的 "收件人电话" 单元格为空，则将电话号码移入正确的单元格。
    """
    try:
        all_rows_cells = _get_unique_row_cells(table)

        name_val_cell = None
        phone_val_cell = None
        for row_cells in all_rows_cells:
            for j, cell in enumerate(row_cells):
                key_text = cell.text.strip()
                if '收件人姓名' in key_text and j + 1 < len(row_cells):
                    name_val_cell = row_cells[j + 1]
                if '收件人电话' in key_text and j + 1 < len(row_cells):
                    phone_val_cell = row_cells[j + 1]

        if name_val_cell is None or phone_val_cell is None:
            return

        name_val = name_val_cell.text.strip()
        phone_val = phone_val_cell.text.strip()

        if phone_val:  # 电话格已有内容，不需要修复
            return

        m = re.search(r'([0-9a-zA-Z\-_+()（）*#@,:：，。一]+)$', name_val)
        if m:
            tail_str = m.group(1)
            if len(re.findall(r'\d', tail_str)) >= 7:
                name_only = name_val[:m.start()]
                _replace_cell_text(name_val_cell, name_only)
                _replace_cell_text(phone_val_cell, tail_str)
                logger.info(f'[fix_name_phone] 修正收件人电话："{name_only}" / "{tail_str}"')
    except Exception as e:
        logger.warning(f"[fix_name_phone] Error: {e}")


def convert_info_docx_multi_page(all_res_list, save_folder, img_name, keep_reference=True):
    """
    将多页的识别结果合并到一个 Word 文档中，最大程度还原原PDF格式
    通过文本框坐标判断同一行的文本，保持原有行结构
    
    all_res_list: 包含 {'img': img, 'res': res, 'page_index': page_index} 的列表
    keep_reference: 是否保留签批区（True时单独保留reference区域图片+标记）
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)

    # 预扫描：判断是否为合同电子签章申请函（相似度>=80%则触发表格后处理）
    _CONTRACT_ESIGN_TITLE = '合同电子签章申请函'
    
    is_contract_esign = False
    for _pd in all_res_list:
        for _r in _pd['res']:
        
            if (_r['type'].lower() == 'title'or _r['type'].lower() == 'header') and _r['res']:
                _title_text = ' '.join(
                    item.get('text', '') if isinstance(item, dict) else str(item)
                    for item in _r['res']
                ).strip()
                
                if SequenceMatcher(None, _title_text, _CONTRACT_ESIGN_TITLE).ratio() >= 0.8:
                    is_contract_esign = True
                    logger.info(f'[contract_esign] 检测到合同电子签章申请函，将启用表格姓名/电话修正')
                    break
        if is_contract_esign:
            break

    for page_idx, page_data in enumerate(all_res_list):
        img = page_data['img']
        res = page_data['res']
        page_index = page_data['page_index']
        
        # 每一页创建新的节（section），这样每页可以有独立的页眉页脚
        if page_idx > 0:
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            # 设置新节的页眉页脚不链接到上一节
            new_section.different_first_page_header_footer = False
            new_section.header.is_linked_to_previous = False
            new_section.footer.is_linked_to_previous = False
        
        # 当前页的section（第一页用默认section，后续页用新创建的section）
        current_section = doc.sections[-1]
        
        # 每页独立的页眉页脚追踪
        header_paragraph = None
        footer_paragraph = None
        
        # 按照垂直位置排序区域（从上到下）
        sorted_res = sorted(res, key=lambda r: (r['bbox'][1], r['bbox'][0]))
        
        # 处理该页的所有内容
        for i, region in enumerate(sorted_res):
            img_idx = region['img_idx']
            region_type = region['type'].lower()
            if len(region['res']) == 0 and region_type != 'reference':
                continue

            if region_type == 'figure':
                excel_save_folder = os.path.join(save_folder, img_name)
                os.makedirs(excel_save_folder, exist_ok=True)
                img_path = os.path.join(excel_save_folder,
                                        '{}_{}.jpg'.format(region['bbox'], img_idx))
                # 如果图片文件不存在，从 region['img'] 保存
                if not os.path.exists(img_path):
                    roi_img = region.get('img')
                    if roi_img is not None and roi_img.size > 0:
                        cv2.imwrite(img_path, roi_img)
                paragraph_pic = doc.add_paragraph()
                paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph_pic.add_run("")
                try:
                    run.add_picture(img_path, width=shared.Inches(5))
                except Exception as e:
                    logger.warning(f"Failed to add picture: {e}")
                    
            elif region_type == 'header':
                # header处理：添加到当前页的页眉
                lines = group_text_by_lines(region['res'])
                
                if header_paragraph is None:
                    # 当前页第一个header，创建页眉段落
                    header = current_section.header
                    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 合并多行为横向拼接
                header_text_parts = []
                for line_items in lines:
                    line_text = merge_line_texts(line_items)
                    header_text_parts.append(line_text)
                
                merged_header_text = '  '.join(header_text_parts)
                
                # 追加到同一段落，用空格分隔
                if header_paragraph.runs:
                    header_paragraph.add_run('  ')  # 段落间分隔
                text_run = header_paragraph.add_run(merged_header_text)
                text_run.font.size = shared.Pt(10)
                        
            elif region_type == 'footer':
                # footer处理：添加到当前页的页脚
                lines = group_text_by_lines(region['res'])
                
                if footer_paragraph is None:
                    # 当前页第一个footer，创建页脚段落
                    footer = current_section.footer
                    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 合并多行为横向拼接
                footer_text_parts = []
                for line_items in lines:
                    line_text = merge_line_texts(line_items)
                    footer_text_parts.append(line_text)
                
                merged_footer_text = '  '.join(footer_text_parts)
                
                # 追加到同一段落，用空格分隔
                if footer_paragraph.runs:
                    footer_paragraph.add_run('  ')  # 段落间分隔
                text_run = footer_paragraph.add_run(merged_footer_text)
                text_run.font.size = shared.Pt(10)
                        
            elif region_type == 'title':
                # title可能有多行，需要分组处理
                lines = group_text_by_lines(region['res'])
                if lines:
                    # 创建一个标题段落
                    title_paragraph = doc.add_heading()
                    
                    # 根据region的x坐标判断对齐方式
                    # 如果左上角在页面左侧30%以内，则左对齐；否则居中
                    bbox = region['bbox']
                    page_width = bbox[2] - bbox[0] if len(bbox) >= 2 else 1000  # 大致估算页面宽度
                    
                    # 使用region的x1坐标（左上角x）
                    x1 = bbox[0]
                    # 假设页面总宽度大约是600（这是典型PDF宽度的相对值）
                    # 如果x1在左侧30%内，则左对齐
                    if x1 < 600 * 0.3:  # 左侧30%
                        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for line_idx, line_items in enumerate(lines):
                        line_text = merge_line_texts(line_items)
                        title_paragraph.add_run(line_text)
                        
                        # 除了最后一行，添加换行符
                        if line_idx < len(lines) - 1:
                            title_paragraph.add_run('\n')
                    
            elif region_type == 'table':
                parser = HtmlToDocx()
                parser.table_style = 'TableGrid'
                try:
                    parser.handle_table(region['res']['html'], doc)
                    if is_contract_esign and doc.tables:
                        _repair_contract_esign_docx_table(doc.tables[-1])
                        _fix_name_phone_in_docx_table(doc.tables[-1])
                except Exception as e:
                    import traceback
                    logger.warning(f"Failed to parse table on page {page_index + 1}, using text fallback: {e}\n{traceback.format_exc()}")
                    print('bebuggggg',region['res']['html'])
                    paragraph = doc.add_paragraph()
                    if isinstance(region['res'], dict) and 'html' in region['res']:
                        html_content = region['res']['html']
                        text_content = re.sub(r'<[^>]+>', '', html_content)
                        paragraph.add_run(text_content)
                    else:
                        paragraph.add_run("[表格识别失败]")
                        
            elif region_type == 'reference':
                # Reference区域处理：始终插入标记，根据keep_reference决定是否插入图片
                
                # 1. 始终插入红底黄字标记
                flag_paragraph = doc.add_paragraph()
                flag_run = flag_paragraph.add_run('检测到签批区')
                flag_run.font.color.rgb = RGBColor(255, 255, 0)  # 黄字
                flag_run.font.bold = True
                flag_run.font.size = shared.Pt(12)
                
                # 设置背景色为红色 (直接使用highlight，不需要shading)
                flag_run.font.highlight_color = 6  # WD_COLOR_INDEX.RED
                
                # 2. 根据keep_reference决定是否插入签批区原始图片
                if keep_reference:
                    excel_save_folder = os.path.join(save_folder, img_name)
                    os.makedirs(excel_save_folder, exist_ok=True)
                    img_path = os.path.join(excel_save_folder,
                                            '{}_{}.jpg'.format(region['bbox'], img_idx))
                    
                    # 如果图片文件不存在，从 region['img'] 保存
                    if not os.path.exists(img_path):
                        roi_img = region.get('img')
                        if roi_img is not None and roi_img.size > 0:
                            cv2.imwrite(img_path, roi_img)
                    
                    # 检查图片文件是否存在
                    if os.path.exists(img_path):
                        paragraph_pic = doc.add_paragraph()
                        paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph_pic.add_run("")
                        try:
                            run.add_picture(img_path, width=shared.Inches(2))
                        except Exception as e:
                            logger.warning(f"Failed to add reference picture: {e}")
                    else:
                        logger.warning(f"Reference region image not found and could not be saved: {img_path}")
                
                # 3. 始终插入reference区域的文字识别结果
                lines = group_text_by_lines(region['res'])
                for line_idx, line_items in enumerate(lines):
                    paragraph = doc.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = shared.Pt(0)
                    paragraph_format.space_before = shared.Pt(0)
                    paragraph_format.line_spacing = 1.0
                    
                    line_text = merge_line_texts(line_items)
                    text_run = paragraph.add_run(line_text)
                    text_run.font.size = shared.Pt(10)
                    
            else:
                # 普通文本区域：按行分组还原
                lines = group_text_by_lines(region['res'])
                
                for line_idx, line_items in enumerate(lines):
                    # 每一行创建一个新段落
                    paragraph = doc.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = shared.Pt(0)  # 减少段落间距
                    paragraph_format.space_before = shared.Pt(0)
                    paragraph_format.line_spacing = 1.0  # 单倍行距
                    
                    # 合并同一行的文本
                    line_text = merge_line_texts(line_items)
                    text_run = paragraph.add_run(line_text)
                    text_run.font.size = shared.Pt(10)

    # save to docx with retry mechanism for locked files
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            # Check if file exists and is locked
            if os.path.exists(docx_path):
                try:
                    # Try to open the file in append mode to check if it's locked
                    with open(docx_path, 'ab') as f:
                        pass
                except (IOError, OSError) as e:
                    # File is locked, create a backup and use a different name
                    logger.warning('File {} is locked by another process, will save as backup'.format(docx_path))
                    timestamp = int(time.time() * 1000) % 100000
                    docx_path_backup = os.path.join(save_folder, '{}_{}_ocr.docx'.format(img_name, timestamp))
                    docx_path = docx_path_backup
                    logger.info('Using backup path: {}'.format(docx_path))
            
            doc.save(docx_path)
            logger.info('docx save to {}'.format(docx_path))
            break  # Success, exit retry loop  
        

             
        except (IOError, OSError, PermissionError) as e:
            if attempt < max_retries - 1:
                logger.warning('Failed to save docx (attempt {}/{}): {}, retrying in 1 second...'.format(
                    attempt + 1, max_retries, str(e)))
                time.sleep(1)
            else:
                # Create backup with timestamp on final failure
                timestamp = int(time.time() * 1000) % 100000
                docx_path_final = os.path.join(save_folder, '{}_{}_ocr.docx'.format(img_name, timestamp))
                try:
                    doc.save(docx_path_final)
                    logger.info('Failed to save to original path, saved to backup: {}'.format(docx_path_final))
                except Exception as final_error:
                    logger.error('Failed to save docx even after retries: {}'.format(str(final_error)))
                    raise
